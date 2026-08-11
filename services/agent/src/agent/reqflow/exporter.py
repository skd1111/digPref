"""reqflow.exporter —— 批次需求文档导出（Markdown / Word）。

文档结构（两种格式一致，spec §6）：
    标题 + 元信息 → 一、批次概览（完成情况统计）→ 二、需求明细（逐条全字段）。
"""

from __future__ import annotations

import io
import time
from typing import Any

from .models import ALL_STATUSES, STATUS_LABEL, ReqBatch, ReqCard

_FEASIBILITY_LABEL = {
    "feasible": "可行",
    "risky": "有风险",
    "infeasible": "不可行",
}


def _feature_refs(card: ReqCard, feature_names: dict[str, str]) -> str:
    """功能点引用渲染：有名字用名字，没有（已删除）标注提示。"""
    if not card.feature_ids:
        return "无"
    parts: list[str] = []
    for fid in card.feature_ids:
        name = feature_names.get(fid)
        parts.append(name if name else f"{fid}（功能点已不存在）")
    return "、".join(parts)


def _overview_lines(batch: ReqBatch, cards: list[ReqCard]) -> list[str]:
    counts = dict.fromkeys(ALL_STATUSES, 0)
    for c in cards:
        if c.status in counts:
            counts[c.status] += 1
    lines = [
        f"- 需求总数：{len(cards)}",
    ]
    for s in ALL_STATUSES:
        lines.append(f"- {STATUS_LABEL[s]}：{counts[s]}")
    return lines


def _card_lines(card: ReqCard, feature_names: dict[str, str]) -> list[str]:
    feasibility = _FEASIBILITY_LABEL.get(card.feasibility, card.feasibility or "未评估")
    lines = [
        f"### {card.id} · {card.title}",
        "",
        f"- 状态：{STATUS_LABEL.get(card.status, card.status)}｜优先级：{card.priority}"
        f"｜版本：v{card.version}",
        f"- 系统名称：{card.system_name}",
        f"- 关联功能点：{_feature_refs(card, feature_names)}",
        f"- 业务价值：{card.business_value or '无'}",
        f"- 改造点：{card.change_points or '无'}",
        f"- 可行性结论：{feasibility}",
        f"- 可行性说明：{card.feasibility_notes or '无'}",
        f"- 对其他功能的影响：{card.impact or '无'}",
        f"- 涉及外部系统：{'、'.join(card.external_systems) if card.external_systems else '无'}",
        "",
    ]
    return lines


def export_markdown(
    batch: ReqBatch,
    cards: list[ReqCard],
    feature_names: dict[str, str],
) -> str:
    """批次导出 Markdown 文档。"""
    cards_sorted = sorted(cards, key=lambda c: c.id)
    lines: list[str] = [
        f"# 需求文档 · {batch.name}",
        "",
        f"- 批次号：{batch.id}",
        f"- 工程：{batch.project_name}",
        f"- 生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## 一、批次概览",
        "",
        *_overview_lines(batch, cards_sorted),
        "",
        "## 二、需求明细",
        "",
    ]
    for card in cards_sorted:
        lines.extend(_card_lines(card, feature_names))
    return "\n".join(lines)


def export_docx(
    batch: ReqBatch,
    cards: list[ReqCard],
    feature_names: dict[str, str],
) -> bytes:
    """批次导出 Word 文档（返回 .docx 字节流）。"""
    from docx import Document  # 局部 import：无 docx 需求时不加载

    cards_sorted = sorted(cards, key=lambda c: c.id)
    doc = Document()
    doc.add_heading(f"需求文档 · {batch.name}", level=0)
    for meta in (
        f"批次号：{batch.id}",
        f"工程：{batch.project_name}",
        f"生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}",
    ):
        doc.add_paragraph(meta)

    doc.add_heading("一、批次概览", level=1)
    for line in _overview_lines(batch, cards_sorted):
        doc.add_paragraph(line.lstrip("- "), style="List Bullet")

    doc.add_heading("二、需求明细", level=1)
    for card in cards_sorted:
        doc.add_heading(f"{card.id} · {card.title}", level=2)
        for line in _card_lines(card, feature_names):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("###"):
                continue  # 标题已用 add_heading
            doc.add_paragraph(stripped.lstrip("- "), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__: list[Any] = ["export_docx", "export_markdown"]
