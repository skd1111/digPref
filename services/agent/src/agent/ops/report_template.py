"""ops 交付物外部报告模板渲染（2026-08-10）。

模板来源：agent/expert_teams/templates.py 解析出的 docx / md 文件。
占位符规范（中文，双花括号）：

  文本类：{{业务名称}} {{专家团}} {{生成时间}} {{材料数量}} {{风险结论}}
  列表类（逐条渲染）：{{材料验收清单}} {{交叉比对清单}} {{问答记录清单}} {{人工确认事项}}
  docx 表格循环行：某表格行内写 {{#材料验收}} … {{/材料验收}}，
    行内单元格可用 {{材料}} {{专家}} {{状态}} {{意见}}，按材料逐行复制。

红线（与 file_to_markdown 同精神）：模板渲染失败只降级，绝不阻塞导出——
调用方拿到 None 后回退内置报告结构。未识别占位符替换为空串并记 warning。
"""

from __future__ import annotations

import logging
import re
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{([^{}]+?)\}\}")
_LOOP_START = "{{#材料验收}}"
_LOOP_END = "{{/材料验收}}"


@dataclass
class TemplateContext:
    """渲染上下文：文本占位符 / 列表占位符 / docx 表格循环行数据。"""

    text: dict[str, str] = field(default_factory=dict)
    lists: dict[str, list[str]] = field(default_factory=dict)
    rows: list[dict[str, str]] = field(default_factory=list)


def _flatten_mapping(ctx: TemplateContext, *, list_joiner: str) -> dict[str, str]:
    """列表占位符按 joiner 拼成字符串（md 渲染用；docx 单独处理换行）。"""
    out = dict(ctx.text)
    for name, lines in ctx.lists.items():
        out[name] = list_joiner.join(lines)
    return out


def _substitute(text: str, mapping: dict[str, str]) -> str:
    """替换 {{占位符}}；未知占位符替换为空并 warning（模板作者可查日志）。"""

    def _repl(m: re.Match[str]) -> str:
        key = m.group(1).strip()
        if key in mapping:
            return mapping[key]
        logger.warning("[report-template] 模板含未识别占位符 {{%s}}，已置空", key)
        return ""

    return _PLACEHOLDER_RE.sub(_repl, text)


def render_markdown_template(template_path: Path, ctx: TemplateContext) -> str:
    """md 模板：纯文本占位符替换（列表逐行）。"""
    raw = template_path.read_text(encoding="utf-8")
    mapping = _flatten_mapping(ctx, list_joiner="\n")
    # 循环标记在 md 里无意义：直接移除，保留行内内容按首条渲染不适用 → 整段删
    raw = raw.replace(_LOOP_START, "").replace(_LOOP_END, "")
    return _substitute(raw, mapping)


def render_docx_template(template_path: Path, ctx: TemplateContext) -> bytes | None:
    """docx 模板：表格循环行复制 + 段落占位符填充。

    失败（依赖缺失/模板损坏）返 None，由调用方降级。
    """
    try:
        import io
        from copy import deepcopy

        from docx import Document
    except Exception as e:
        logger.warning("[report-template] python-docx unavailable: %s", e)
        return None

    try:
        doc = Document(str(template_path))
        text_map = dict(ctx.text)

        # 1) 表格循环行：{{#材料验收}} 所在行按 rows 逐条复制
        for table in doc.tables:
            _expand_loop_rows(table, ctx.rows, deepcopy)

        # 2) 全部段落占位符填充（正文 + 表格单元格）
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        for p in paragraphs:
            _fill_paragraph(p, text_map, ctx.lists)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.warning("[report-template] render %s failed: %s", template_path.name, e)
        return None


def _expand_loop_rows(
    table: Any, rows: list[dict[str, str]], deepcopy: Callable[[Any], Any]
) -> None:
    """把含 {{#材料验收}} 的模板行替换为每份材料一行（倒序插入保持顺序）。"""
    template_row = None
    for row in table.rows:
        full = "".join(cell.text for cell in row.cells)
        if _LOOP_START in full:
            template_row = row
            break
    if template_row is None:
        return
    template_tr = template_row._tr
    cell_texts = [cell.text for cell in template_row.cells]
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for item in reversed(rows):
        new_tr = deepcopy(template_tr)
        template_tr.addprevious(new_tr)
        for cell, tmpl_text in zip(new_tr.findall(f"{w_ns}tc"), cell_texts):
            # 直接重写单元格文本：循环行样式以模板行为准，占位符填充优先
            value = _substitute(tmpl_text.replace(_LOOP_START, "").replace(_LOOP_END, ""), item)
            _set_tc_text(cell, value.strip())
    template_tr.getparent().remove(template_tr)


def _set_tc_text(tc: Any, value: str) -> None:
    """重写 <w:tc> 单元格文本（保留首个段落结构）。"""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs = tc.findall(f"{ns}p")
    if not paragraphs:
        return
    first = paragraphs[0]
    for extra in paragraphs[1:]:
        tc.remove(extra)
    runs = first.findall(f"{ns}r")
    if runs:
        for extra in runs[1:]:
            first.remove(extra)
        texts = runs[0].findall(f"{ns}t")
        if texts:
            texts[0].text = value
            for extra in texts[1:]:
                runs[0].remove(extra)
        else:
            t = first.makeelement(f"{ns}t", {})
            t.text = value
            runs[0].append(t)


def _fill_paragraph(p: Any, text_map: dict[str, str], lists: dict[str, list[str]]) -> None:
    """段落占位符填充：合并 runs 文本 → 替换；列表占位符逐条换行。

    Word 常把 {{占位符}} 拆进多个 run，故先合并再写回首 run（牺牲 run 级
    局部样式，保证占位符一定命中——模板场景下可接受）。
    """
    if not p.runs:
        return
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return

    # 切分为 文本段 / 列表占位符 序列
    segments: list[tuple[str, str]] = []  # (kind, value) kind ∈ {'text', 'list'}
    cursor = 0
    pattern = re.compile("|".join(re.escape("{{" + n + "}}") for n in lists)) if lists else None
    if pattern:
        for m in pattern.finditer(full):
            if m.start() > cursor:
                segments.append(("text", full[cursor : m.start()]))
            segments.append(("list", m.group(0)[2:-2]))
            cursor = m.end()
    if cursor < len(full):
        segments.append(("text", full[cursor:]))

    first_run = p.runs[0]
    for r in p.runs[1:]:
        r.text = ""
    wrote = False
    for kind, value in segments:
        if kind == "text":
            replaced = _substitute(value, text_map)
            if replaced:
                if not wrote:
                    first_run.text = replaced
                    wrote = True
                else:
                    first_run.add_text(replaced)
        else:
            for line in lists.get(value, []):
                if wrote:
                    first_run.add_break()
                first_run.add_text(line)
                wrote = True
    if not wrote:
        first_run.text = ""
