"""ops.api —— FastAPI 路由（运营工作台业务记录，Phase 2H）。

端点：
  POST   /ops/records               — 创建业务记录卡片
  GET    /ops/records               — 列表（feature_id / project_name / limit 过滤）
  GET    /ops/records/{id}          — 详情
  DELETE /ops/records/{id}          — 删除
  POST   /ops/records/summarize     — AI 根据会话 + 功能点 + Skill 生成总结草稿

专家验收工作流（2026-08-10，取代运营模式的传统大 Chat）：
  GET    /ops/case                  — 获取 Case（材料文件 + 问答记录 + 交付草稿）
  POST   /ops/case/files            — 上传材料给对应专家
  POST   /ops/case/files/{id}/review   — AI 专家审核验收
  POST   /ops/case/files/{id}/override — 人工改判
  DELETE /ops/case/files/{id}       — 删除材料
  POST   /ops/case/ask              — 向专家提问（模板/清单类请求 → 生成直填草稿）
  POST   /ops/case/export           — 全部验收后打包导出交付物 zip

交付草稿（BUGFIX #78：要模板不给一大段文字，界面直填）：
  GET    /ops/case/drafts           — 草稿列表
  PUT    /ops/case/drafts/{id}      — 保存填写值
  POST   /ops/case/drafts/{id}/submit — 提交：渲染成材料自动审核，通过即入交付物

LLM 优先级与 biznav/reqflow 一致：本地 Ollama → DB 内网 → DB 云端（三级降级链）。
"""

from __future__ import annotations

import json
import logging
import os
import sys
import time
import zipfile
from pathlib import Path
from typing import Any

from fastapi import APIRouter, HTTPException, Query
from pydantic import BaseModel, Field

from .cases import (
    ALL_FILE_STATUSES,
    FILE_PASSED,
    FILE_PENDING,
    FILE_REJECTED,
    FILE_REVIEWING,
    CaseStorage,
    dump_json,
    is_image_file,
    make_case_id,
)
from .models import ALL_RESULTS
from .storage import BusinessRecordStorage

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ops", tags=["ops"])


# ---------------------------------------------------------------------------
# Storage 单例
# ---------------------------------------------------------------------------


def _default_db_path() -> str:
    appdata = os.environ.get("APPDATA")
    if appdata:
        return os.path.join(appdata, "eaide", "ops.db")
    return os.path.expanduser("~/.eaide/ops.db")


_storage: BusinessRecordStorage | None = None


def _get_storage(db_path: str | None = None) -> BusinessRecordStorage:
    global _storage
    if db_path:
        return BusinessRecordStorage(db_path)
    if _storage is None:
        _storage = BusinessRecordStorage(os.environ.get("EAIDE_OPS_DB", _default_db_path()))
    return _storage


def _reset_storage_for_tests() -> None:
    global _storage
    _storage = None


_case_storage: CaseStorage | None = None


def _get_case_storage(db_path: str | None = None) -> CaseStorage:
    global _case_storage
    if db_path:
        return CaseStorage(db_path)
    if _case_storage is None:
        _case_storage = CaseStorage(os.environ.get("EAIDE_OPS_DB", _default_db_path()))
    return _case_storage


def _reset_case_storage_for_tests() -> None:
    global _case_storage
    _case_storage = None


# ---------------------------------------------------------------------------
# LLM 降级链（与 biznav/reqflow 同一形态）
# ---------------------------------------------------------------------------


def _make_summarize_llm():
    """构造 async (messages) -> str；三级全失败抛 RuntimeError。"""

    async def _call(backend, messages: list[dict]) -> str:
        return str(await backend.extract_chat(messages) or "")

    async def _client(messages: list[dict]) -> str:
        from agent.llm.router import LMRouter

        router = LMRouter()
        if router._mock_mode:
            return await _call(router.mock, messages)

        failures: list[str] = []
        # 1/3 本地 Ollama（本地优先）
        try:
            text = await _call(router.ollama, messages)
            if text.strip():
                return text
            failures.append("本地 Ollama 返回空")
        except Exception as e:
            failures.append(f"本地 Ollama 不可用: {e}")
        # 2/3 内网
        try:
            private = await router._build_private_client()
            if private is not None:
                try:
                    text = await _call(private, messages)
                    if text.strip():
                        return text
                    failures.append("内网模型返回空")
                except Exception as e:
                    failures.append(f"内网模型不可用: {e}")
            else:
                failures.append("内网模型未启用")
        except Exception as e:
            failures.append(f"内网模型查询失败: {e}")
        # 3/3 云端
        try:
            cloud = await router._build_cloud_client()
            if cloud is not None:
                text = await _call(cloud, messages)
                if text.strip():
                    return text
                failures.append("云端模型返回空")
            else:
                failures.append("云端模型未配置")
        except Exception as e:
            failures.append(f"云端模型不可用: {e}")

        raise RuntimeError(
            "所有 LLM 后端均不可用（" + "；".join(failures) + "）。"
            "请启动本地 Ollama，或在「模型管理」中配置可用的云端/内网模型"
        )

    return _client


# ---------------------------------------------------------------------------
# Pydantic 模型
# ---------------------------------------------------------------------------


class CreateRecordRequest(BaseModel):
    project_name: str = ""
    feature_id: str = ""
    business_type: str = ""
    title: str = ""
    summary: str = ""
    materials_checked: list[str] = Field(default_factory=list)
    materials_missing: list[str] = Field(default_factory=list)
    risk_points: list[str] = Field(default_factory=list)
    result: str = "done"
    skill_id: str = ""
    session_id: str = ""
    source: str = "ai"
    created_by: str = ""


class SummarizeRequest(BaseModel):
    feature_id: str = ""
    project_name: str = ""
    business_type: str = ""
    conversation: list[dict[str, str]] = Field(default_factory=list)
    session_id: str = ""


def _validate_result(result: str) -> str:
    if result not in ALL_RESULTS:
        raise HTTPException(400, f"invalid result: {result} (allowed: {', '.join(ALL_RESULTS)})")
    return result


def _load_feature(feature_id: str, project_name: str) -> dict[str, Any] | None:
    if not feature_id:
        return None
    try:
        from agent.biznav.api import _get_storage as _biznav_storage

        f = _biznav_storage().get(feature_id, project_name)
    except Exception as e:
        logger.warning("[ops] biznav storage unavailable: %s", e)
        return None
    if f is None:
        return None
    return {
        "id": f.id,
        "name": f.name,
        "description": f.description,
        "category": f.category,
        "business_rules": [r.text for r in f.business_rules],
        "skill_id": f.skill_id,
    }


def _load_skill(skill_id: str) -> dict[str, Any] | None:
    if not skill_id:
        return None
    try:
        from agent.skills.api import _loader

        s = _loader.get(skill_id)
    except Exception as e:
        logger.warning("[ops] skills loader unavailable: %s", e)
        return None
    if s is None:
        return None
    return {
        "id": s.id,
        "name": s.name,
        "description": s.description,
        "system_prompt": s.system_prompt,
        "trigger_keywords": list(s.trigger_keywords),
    }


def _render_conversation(conversation: list[dict[str, str]]) -> str:
    lines: list[str] = []
    for m in conversation[-30:]:
        role = str(m.get("role", ""))
        content = str(m.get("content", "")).strip()
        if not content:
            continue
        lines.append(f"[{role}] {content[:1200]}")
    return "\n".join(lines)


def _build_summarize_prompt(
    *,
    feature: dict[str, Any] | None,
    skill: dict[str, Any] | None,
    business_type: str,
    conversation_text: str,
) -> list[dict[str, str]]:
    sys_ctx: list[str] = [
        "你是银行运营一线的工作台 AI。用户完成一笔业务后，请把这段会话总结成一张"
        "结构化「业务记录卡片」，要求：",
        "1. 只基于会话与提供的功能点/Skill 经验，不要编造；",
        "2. 材料核对：区分「已具备材料」与「缺失/待补材料」；",
        "3. 风险点：只列会话中出现的风险或 Skill 经验中提示的高风险项；",
        "4. 输出严格 JSON（不要 Markdown 代码块），字段：",
        '   {"title": str, "business_type": str, "summary": str, '
        '"materials_checked": [str], "materials_missing": [str], '
        '"risk_points": [str], "result": "done|pending|rejected|follow_up"}',
        "5. 如果会话不足以判断，result 用 pending，缺失材料写进 materials_missing。",
    ]
    if feature:
        sys_ctx.append(f"【功能点】{feature['name']}（{feature['id']}）：{feature['description']}")
        if feature.get("business_rules"):
            sys_ctx.append("业务规则：\n" + "\n".join(f"- {r}" for r in feature["business_rules"]))
    if skill:
        sys_ctx.append(f"【Skill 经验】{skill['name']}：\n{skill['system_prompt'][:3000]}")
    return [
        {"role": "system", "content": "\n".join(sys_ctx)},
        {
            "role": "user",
            "content": (
                f"业务类型：{business_type or '未指定'}\n\n"
                f"会话记录：\n{conversation_text or '（无会话记录）'}\n\n"
                "请输出业务记录卡片 JSON。"
            ),
        },
    ]


async def _audit(event: str, payload: dict[str, Any]) -> None:
    try:
        from agent.audit.store import audit

        await audit(event, payload)
    except Exception as e:
        logger.warning("[ops] audit emit failed: %s", e)


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@router.post("/records")
async def create_record(req: CreateRecordRequest) -> dict:
    storage = _get_storage()
    _validate_result(req.result)
    rec = storage.create(
        project_name=req.project_name,
        feature_id=req.feature_id,
        business_type=req.business_type,
        title=req.title,
        summary=req.summary,
        materials_checked=req.materials_checked,
        materials_missing=req.materials_missing,
        risk_points=req.risk_points,
        result=req.result,
        skill_id=req.skill_id,
        session_id=req.session_id,
        source=req.source,
        created_by=req.created_by,
    )
    await _audit(
        "ops_record_create",
        {"record_id": rec.id, "feature_id": rec.feature_id, "title": rec.title},
    )
    return rec.to_dict()


@router.get("/records")
async def list_records(
    feature_id: str | None = Query(None),
    project_name: str | None = Query(None),
    limit: int = Query(100),
) -> dict:
    storage = _get_storage()
    records = storage.list(feature_id=feature_id, project_name=project_name, limit=limit)
    return {
        "records": [r.to_dict() for r in records],
        "total": len(records),
    }


@router.get("/records/{record_id}")
async def get_record(record_id: str) -> dict:
    storage = _get_storage()
    rec = storage.get(record_id)
    if rec is None:
        raise HTTPException(404, f"record {record_id} not found")
    return rec.to_dict()


@router.delete("/records/{record_id}")
async def delete_record(record_id: str) -> dict:
    storage = _get_storage()
    storage.delete(record_id)
    await _audit("ops_record_delete", {"record_id": record_id})
    return {"ok": True, "record_id": record_id}


@router.post("/records/summarize")
async def summarize_record(req: SummarizeRequest) -> dict:
    feature = _load_feature(req.feature_id, req.project_name)
    skill_id = feature.get("skill_id") if feature else ""
    skill = _load_skill(skill_id or "")
    conversation_text = _render_conversation(req.conversation)
    messages = _build_summarize_prompt(
        feature=feature,
        skill=skill,
        business_type=req.business_type,
        conversation_text=conversation_text,
    )
    llm_call = _make_summarize_llm()
    try:
        text = await llm_call(messages)
    except RuntimeError as e:
        raise HTTPException(502, str(e)) from e

    draft: dict[str, Any] = {}
    try:
        text = text.strip()
        if text.startswith("```"):
            text = text.strip("`")
            if text.startswith("json"):
                text = text[4:]
        draft = json.loads(text)
    except (ValueError, TypeError) as e:
        logger.warning("[ops] summarize returned non-JSON (%s): %s", e, text[:200])
        raise HTTPException(
            502,
            "模型未返回有效 JSON，请重试或手动填写业务记录",
        ) from e

    if not isinstance(draft, dict):
        raise HTTPException(502, "模型返回格式错误（非对象）")
    if not str(draft.get("title", "")).strip():
        draft["title"] = req.business_type or feature.get("name", "") if feature else ""
    result = str(draft.get("result", "pending"))
    if result not in ALL_RESULTS:
        result = "pending"
    return {
        "draft": {
            "title": str(draft.get("title", "")),
            "business_type": str(draft.get("business_type", req.business_type)),
            "summary": str(draft.get("summary", "")),
            "materials_checked": [str(x) for x in (draft.get("materials_checked") or [])],
            "materials_missing": [str(x) for x in (draft.get("materials_missing") or [])],
            "risk_points": [str(x) for x in (draft.get("risk_points") or [])],
            "result": result,
        }
    }


# ---------------------------------------------------------------------------
# 专家验收工作流（2026-08-10）
# ---------------------------------------------------------------------------


class CaseFileAddRequest(BaseModel):
    case_id: str = Field(min_length=1, max_length=512)
    team_id: str = ""
    member_key: str = Field(min_length=1, max_length=128)
    file_name: str = Field(min_length=1, max_length=512)
    content_base64: str = Field(min_length=1)


class CaseFileOverrideRequest(BaseModel):
    status: str
    note: str = ""


class CaseAskRequest(BaseModel):
    case_id: str = Field(min_length=1, max_length=512)
    team_id: str = ""
    member_key: str = Field(min_length=1, max_length=128)
    question: str = Field(min_length=1, max_length=4000)


class DraftSaveRequest(BaseModel):
    """保存草稿填写值（BUGFIX #78：界面直填）。"""

    values: dict[str, Any] = Field(default_factory=dict)


class CaseExportRequest(BaseModel):
    case_id: str = Field(min_length=1, max_length=512)
    target_path: str = Field(min_length=1, max_length=4096)
    feature_name: str = ""
    team_id: str = ""  # 专家团 id（2026-08-10）：用于定位外部报告模板
    team_name: str = ""
    checklist: list[str] = Field(default_factory=list)


def _resolve_member(team_id: str, member_key: str) -> dict[str, Any] | None:
    """按 team_id + 成员名解析专家人设（prompt / focus_points / outputs）。"""
    if not team_id or not member_key:
        return None
    try:
        from agent.expert_teams.api import get_loader

        team = get_loader().get(team_id)
    except Exception as e:
        logger.warning("[ops-case] expert teams loader unavailable: %s", e)
        return None
    if team is None:
        return None
    for m in team.members:
        if m.name == member_key:
            return m.to_dict()
    return None


def _extract_material_text(file_path: str, file_name: str) -> str:
    """提取非图片类材料的文本（解析失败返空串，不阻断审核）。"""
    try:
        from agent.doc_review.parser import parse_document

        parsed = parse_document(file_path)
        return str(getattr(parsed, "full_text", "") or "")[:20000]
    except Exception as e:
        logger.warning("[ops-case] parse %s failed: %s", file_name, e)
        return ""


async def _vision_describe(file_path: str) -> str:
    """本地视觉模型描述图片内容（不可用返空串）。"""
    try:
        from agent.config import settings
        from agent.llm.local_vision import LocalVisionClient

        client = LocalVisionClient(
            base_url=settings.local_vision_base_url or "http://127.0.0.1:8082/v1",
            model=settings.local_vision_model or "moondream2",
        )
        image_bytes = _read_file_bytes(file_path)
        return str(
            await client.understand_screenshot(
                image_bytes,
                prompt="请详细描述这份材料图片中的关键信息（名称、编号、日期、公章等）。",
            )
            or ""
        )
    except Exception as e:
        logger.debug("[ops-case] local vision unavailable: %s", e)
        return ""


def _read_file_bytes(path: str) -> bytes:
    """同步读文件（单独抽出，避免 async 函数内直接阻塞 IO 的 lint 告警）。"""
    return Path(path).read_bytes()


def _path_parent_exists(target: Path) -> bool:
    return target.parent.exists()


_REVIEW_OUTPUT_RULE = (
    "请验收这份材料，输出严格 JSON（不要 Markdown 代码块）："
    '{"verdict": "passed" 或 "rejected", "note": "验收意见（通过依据或打回原因，中文，不超 200 字）", '
    '"fields": [{"field": "要素名", "value": "要素值", "confidence": 0到1的小数}], '
    '"evidence": ["支撑结论的原文摘录"], '
    '"reject_spans": [{"quote": "问题所在的原文片段", "advice": "该怎么改（不超 60 字）"}]}。'
    "fields 要求：提取材料中的关键要素（如企业名称/统一社会信用代码/法人/有效期/"
    "金额/日期等，按材料类型自行判断），找不到确切值但应有该要素时 confidence 给低于 0.6；"
    "evidence 要求：直接引用材料原文片段（最多 3 条，每条不超 80 字），严禁编造；"
    "reject_spans 要求：仅在 verdict=rejected 时填写，逐条指出问题在原文中的具体位置（"
    "quote 必须是材料里真实存在的片段，最多 5 条）；passed 时给空数组；"
    "判定标准：材料内容与你的职责相关、信息可辨认且无明显矛盾 → passed；"
    "内容缺失、与职责无关、关键信息不可辨认或存在明显疑点 → rejected。"
)

# 规则兑底提取（LLM 不可用/未返回 fields 时）：正则命中 = 高置信（0.95）
_FIELD_PATTERNS: tuple[tuple[str, str], ...] = (
    ("统一社会信用代码", r"[0-9A-HJ-NPQRTUWXY]{2}\d{6}[0-9A-HJ-NPQRTUWXY]{10}"),
    ("身份证号", r"\d{17}[\dXx]"),
    ("手机号", r"1[3-9]\d{9}"),
    ("日期", r"\d{4}\s*[\-/年]\s*\d{1,2}\s*[\-/月]\s*\d{1,2}\s*日?"),
    ("金额", r"[\d,]+(?:\.\d+)?\s*(?:万元|亿元|元)"),
)


def _extract_fields_regex(text: str, limit: int = 12) -> list[dict]:
    """正则提取关键要素（纠错成本陷阱的兑底：LLM 抽风时员工仍能看到结构化要素）。"""
    import re

    out: list[dict] = []
    for field, pattern in _FIELD_PATTERNS:
        seen: set[str] = set()
        for m in re.finditer(pattern, text):
            value = m.group(0).strip()
            if value in seen:
                continue
            seen.add(value)
            out.append({"field": field, "value": value, "confidence": 0.95})
            if len(out) >= limit:
                return out
    return out


def _build_review_messages(
    member: dict[str, Any] | None, file_name: str, material_text: str
) -> list[dict[str, str]]:
    persona = (
        member.get("prompt") or f"你是{member.get('name', '审核专家')}，{member.get('role', '')}"
        if member
        else "你是银行运营审核专家。"
    )
    focus = member.get("focus_points") or [] if member else []
    sys_parts = [
        persona,
        "当前场景：客户经理正在办理业务，向你上传了一份待验收材料，请你按自己的职责审核验收。",
    ]
    if focus:
        sys_parts.append("你的关注点：\n" + "\n".join(f"- {x}" for x in focus))
    return [
        {"role": "system", "content": "\n\n".join(sys_parts)},
        {
            "role": "user",
            "content": (
                f"材料文件名：{file_name}\n\n材料内容：\n{material_text or '（无可提取文本）'}\n\n"
                + _REVIEW_OUTPUT_RULE
            ),
        },
    ]


def _parse_review_result(
    text: str, source_text: str
) -> tuple[str, str, list[dict], list[str], list[dict]]:
    """解析 LLM 验收结果 → (status, note, fields, evidence, reject_marks)。

    - JSON 解析失败 → 保守打回 + 正则兑底要素；
    - LLM fields 与正则 fields 去重合并（LLM 优先）；
    - evidence 只保留真实存在于原文的摘录（容忍空白差异，防幻觉证据链）；
    - reject_marks 同样白名单校验（quote 必须在原文中），模型未返回时宽容置空（BUGFIX #80）。
    """
    fields = _extract_fields_regex(source_text)
    evidence: list[str] = []
    marks: list[dict] = []
    t = (text or "").strip()
    if t.startswith("```"):
        t = t.strip("`")
        if t.startswith("json"):
            t = t[4:]
    try:
        data = json.loads(t)
        if isinstance(data, dict):
            verdict = str(data.get("verdict", "")).lower()
            note = str(data.get("note", "")).strip()
            llm_fields = data.get("fields")
            if isinstance(llm_fields, list) and llm_fields:
                merged: list[dict] = []
                seen_keys: set[str] = set()
                for f in llm_fields:
                    if not isinstance(f, dict):
                        continue
                    fname = str(f.get("field", "")).strip()
                    fvalue = str(f.get("value", "")).strip()
                    if not fname or not fvalue:
                        continue
                    key = f"{fname}::{fvalue}"
                    if key in seen_keys:
                        continue
                    seen_keys.add(key)
                    try:
                        conf = float(f.get("confidence", 0.8))
                    except (TypeError, ValueError):
                        conf = 0.8
                    merged.append(
                        {"field": fname, "value": fvalue, "confidence": max(0.0, min(1.0, conf))}
                    )
                for rf in fields:
                    if f"{rf['field']}::{rf['value']}" not in seen_keys:
                        merged.append(rf)
                fields = merged[:15]
            ev = data.get("evidence")
            if isinstance(ev, list) and source_text:
                norm_source = "".join(source_text.split())
                for e in ev[:3]:
                    s = str(e).strip()
                    if s and "".join(s.split()) in norm_source:
                        evidence.append(s[:120])
            # 打回定位（BUGFIX #80）：quote 必须在原文中真实存在，否则丢弃（防幻觉标注）
            spans = data.get("reject_spans")
            if isinstance(spans, list) and source_text:
                norm_src = "".join(source_text.split())
                for sp in spans[:5]:
                    if not isinstance(sp, dict):
                        continue
                    quote = str(sp.get("quote", "") or "").strip()
                    advice = str(sp.get("advice", "") or "").strip()
                    if not quote or "".join(quote.split()) not in norm_src:
                        continue
                    marks.append({"quote": quote[:120], "advice": advice[:200]})
            if verdict in ("passed", "pass", "ok", "通过"):
                return FILE_PASSED, note or "AI 审核通过", fields, evidence, []
            if verdict in ("rejected", "reject", "fail", "打回"):
                return FILE_REJECTED, note or "AI 审核未通过", fields, evidence, marks
    except (ValueError, TypeError):
        pass
    return (
        FILE_REJECTED,
        f"AI 审核结论无法解析，已保守打回。原文：{t[:300]}",
        fields,
        evidence,
        [],
    )


def _normalize_file_row(row: dict | None) -> dict:
    """存储行 → API 响应：extracted_fields/evidence/reject_marks 从 JSON 文本解析为数组。"""
    if row is None:
        return {}
    out = dict(row)
    for key in ("extracted_fields", "evidence", "reject_marks"):
        val = out.get(key)
        if isinstance(val, str):
            try:
                out[key] = json.loads(val or "[]")
            except (ValueError, TypeError):
                out[key] = []
        elif val is None:
            out[key] = []
    return out


@router.get("/case")
async def get_case(
    project_name: str = Query(""),
    feature_id: str = Query(""),
    case_id: str = Query(""),
) -> dict:
    cid = case_id or make_case_id(project_name, feature_id)
    store = _get_case_storage()
    return {
        "case_id": cid,
        "files": [_normalize_file_row(f) for f in store.list_files(cid)],
        "qa": store.list_qa(cid),
        "drafts": [_normalize_draft_row(d) for d in store.list_drafts(cid)],
    }


@router.delete("/case")
async def clear_case(
    project_name: str = Query(""),
    feature_id: str = Query(""),
    case_id: str = Query(""),
) -> dict:
    """清空 Case（BUGFIX #85）：重新开始办理 —— 材料/问答/草稿/落盘文件全部清除。"""
    cid = case_id or make_case_id(project_name, feature_id)
    store = _get_case_storage()
    counts = store.clear_case(cid)
    await _audit("ops_case_clear", {"case_id": cid, **counts})
    return {"ok": True, "case_id": cid, **counts}


@router.post("/case/files")
async def case_add_file(req: CaseFileAddRequest) -> dict:
    store = _get_case_storage()
    try:
        path = store.save_upload(req.case_id, req.file_name, req.content_base64)
    except Exception as e:
        raise HTTPException(400, f"文件保存失败: {e}") from e
    row = store.add_file(
        case_id=req.case_id,
        team_id=req.team_id,
        member_key=req.member_key,
        file_name=req.file_name,
        file_path=str(path),
    )
    await _audit(
        "ops_case_file_add",
        {
            "case_id": req.case_id,
            "file_id": row["id"],
            "member_key": req.member_key,
            "file_name": req.file_name,
        },
    )
    return _normalize_file_row(row)


async def _do_review_file(store: CaseStorage, file_id: str) -> dict:
    """AI 专家审核单个材料（端点与草稿提交共用，BUGFIX #78）。"""
    row = store.get_file(file_id)
    if row is None:
        raise HTTPException(404, f"case file {file_id} not found")

    member = _resolve_member(str(row["team_id"]), str(row["member_key"]))
    store.update_file(file_id, status=FILE_REVIEWING)

    # 图片类：优先本地 vision；不可用 → 降级为关注点核对清单 + 待人工确认
    source_text = ""
    if is_image_file(str(row["file_name"])):
        description = await _vision_describe(str(row["file_path"]))
        if description.strip():
            source_text = description
            messages = _build_review_messages(member, str(row["file_name"]), description)
        else:
            focus = (member or {}).get("focus_points") or []
            note = "本地视觉模型不可用，无法自动识别图片内容。请人工按以下要点核对后改判：\n" + (
                "\n".join(f"- {x}" for x in focus) if focus else "- 材料真实、清晰、与业务相关"
            )
            # 低置信占位要素：前端标红提醒「未能自动提取，需人工核对」（纠错成本陷阱的诚实表达）
            placeholder_fields = [
                {
                    "field": "图片内容",
                    "value": "本地视觉模型不可用，未能自动提取关键要素",
                    "confidence": 0.1,
                }
            ]
            updated = store.update_file(
                file_id,
                status=FILE_PENDING,
                review_note=note,
                reviewed_by="ai",
                extracted_fields=dump_json(placeholder_fields),
                evidence="[]",
            )
            await _audit(
                "ops_case_file_review",
                {"file_id": file_id, "verdict": FILE_PENDING, "mode": "vision_fallback"},
            )
            return _normalize_file_row(updated)
    else:
        source_text = _extract_material_text(str(row["file_path"]), str(row["file_name"]))
        messages = _build_review_messages(member, str(row["file_name"]), source_text)

    llm_call = _make_summarize_llm()
    try:
        raw = await llm_call(messages)
    except RuntimeError as e:
        # 三级 LLM 均不可用 → 回退待审核状态，但正则要素照样给（员工不必白等）
        store.update_file(
            file_id,
            status=FILE_PENDING,
            review_note=f"AI 审核不可用：{e}",
            reviewed_by="",
            extracted_fields=dump_json(_extract_fields_regex(source_text)),
        )
        raise HTTPException(502, str(e)) from e

    status, note, fields, evidence, marks = _parse_review_result(raw, source_text)
    updated = store.update_file(
        file_id,
        status=status,
        review_note=note,
        reviewed_by="ai",
        extracted_fields=dump_json(fields),
        evidence=dump_json(evidence),
        # 通过时清空旧打回定位；打回时写入新定位（BUGFIX #80）
        reject_marks=dump_json(marks),
    )
    await _audit(
        "ops_case_file_review",
        {"file_id": file_id, "member_key": row["member_key"], "verdict": status},
    )
    return _normalize_file_row(updated)


@router.post("/case/files/{file_id}/review")
async def case_review_file(file_id: str) -> dict:
    return await _do_review_file(_get_case_storage(), file_id)


@router.get("/case/files/{file_id}/content")
async def case_file_content(file_id: str) -> dict:
    """交付物柜预览（BUGFIX #79）：返回材料文件 base64 内容。"""
    import base64

    store = _get_case_storage()
    row = store.get_file(file_id)
    if row is None:
        raise HTTPException(404, f"case file {file_id} not found")
    path = Path(str(row.get("file_path", "")))
    if not path.is_file():
        raise HTTPException(404, "文件已不存在（可能被清理）")
    raw = path.read_bytes()
    return {
        "file_name": str(row.get("file_name", "")),
        "content_base64": base64.b64encode(raw).decode("ascii"),
    }


class CaseFileSaveAsRequest(BaseModel):
    target_path: str = Field(min_length=1, max_length=4096)


@router.post("/case/files/{file_id}/save-as")
async def case_file_save_as(file_id: str, req: CaseFileSaveAsRequest) -> dict:
    """交付物柜另存（BUGFIX #79）：后端直接复制文件，避免大体积 base64 过桥。"""
    import shutil

    store = _get_case_storage()
    row = store.get_file(file_id)
    if row is None:
        raise HTTPException(404, f"case file {file_id} not found")
    src = Path(str(row.get("file_path", "")))
    if not src.is_file():
        raise HTTPException(404, "文件已不存在（可能被清理）")
    target = Path(req.target_path)
    if not _path_parent_exists(target):
        raise HTTPException(400, f"保存目录不存在: {target.parent}")
    shutil.copyfile(src, target)
    await _audit(
        "ops_case_file_save_as",
        {"file_id": file_id, "target_path": str(target)},
    )
    return {"ok": True, "path": str(target)}


@router.post("/case/files/{file_id}/override")
async def case_override_file(file_id: str, req: CaseFileOverrideRequest) -> dict:
    if req.status not in ALL_FILE_STATUSES:
        raise HTTPException(400, f"invalid status: {req.status}")
    store = _get_case_storage()
    before = store.get_file(file_id)
    if before is None:
        raise HTTPException(404, f"case file {file_id} not found")
    updated = store.update_file(
        file_id,
        status=req.status,
        review_note=req.note,
        reviewed_by="human",
    )
    # 纠错闭环（铁律 2，2026-08-10）：AI 结论被人工改判 → 落纠错样本，
    # 让 AI 越用越聪明，而不是永远犯同样的错
    if str(before.get("reviewed_by", "")) == "ai" and str(before.get("status", "")) != req.status:
        store.add_correction(
            case_id=str(before.get("case_id", "")),
            file_id=file_id,
            member_key=str(before.get("member_key", "")),
            file_name=str(before.get("file_name", "")),
            ai_status=str(before.get("status", "")),
            ai_note=str(before.get("review_note", "")),
            human_status=req.status,
            human_note=req.note,
        )
        await _audit(
            "ops_case_correction",
            {
                "file_id": file_id,
                "ai_status": str(before.get("status", "")),
                "human_status": req.status,
            },
        )
    await _audit(
        "ops_case_file_override",
        {"file_id": file_id, "status": req.status, "note": req.note},
    )
    return _normalize_file_row(updated)


@router.get("/case/corrections")
async def case_list_corrections(case_id: str = Query("")) -> dict:
    """人工纠错样本列表（供后续提示词/模型改进分析，不指定 case_id 返全部）。"""
    store = _get_case_storage()
    rows = store.list_corrections(case_id or None)
    return {"corrections": rows, "total": len(rows)}


@router.delete("/case/files/{file_id}")
async def case_delete_file(file_id: str) -> dict:
    store = _get_case_storage()
    row = store.delete_file(file_id)
    if row is None:
        raise HTTPException(404, f"case file {file_id} not found")
    await _audit(
        "ops_case_file_delete",
        {"file_id": file_id, "file_name": row.get("file_name", "")},
    )
    return {"ok": True, "file_id": file_id}


def _knowledge_base_dir() -> Path | None:
    """仓库内 knowledge-base 目录（合规/法律/数据安全/资金风险/案例库 md）。

    支持 EAIDE_KB_DIR 环境变量覆盖（部署形态不同时指向实际目录）。
    """
    override = os.environ.get("EAIDE_KB_DIR")
    if override:
        p = Path(override)
        return p if p.is_dir() else None
    candidate = Path(__file__).resolve().parents[5] / "knowledge-base"
    if candidate.is_dir():
        return candidate
    # PyInstaller 单文件打包后源码在 _MEIPASS 解压目录，parents[5] 上溯不到；
    # 回退到 spec datas 打进来的内置副本（与 doc_review/knowledge.py 同策略）。
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        bundled = Path(meipass) / "knowledge-base"
        if bundled.is_dir():
            return bundled
    return None


_KB_STOPWORDS = frozenset(
    {"请问", "什么", "怎么", "如何", "哪些", "需要", "注意", "问题", "办理", "业务", "材料", "专家"}
)


def _search_local_knowledge(question: str, limit: int = 3) -> list[dict]:
    """知识库关键词检索（找制度变问助手，2026-08-10）。

    对 knowledge-base/*.md 按关键词命中计分，返回带出处的片段：
    [{title, source, snippet}]。知识库不存在/无命中 → 空列表（降级不阻塞）。
    """
    kb_dir = _knowledge_base_dir()
    if kb_dir is None:
        return []
    # 关键词：中文无分词 → 对连续中文串取 2/3 字 n-gram（含停用词的丢弃），
    # 再并上英文/数字词；命中计分仅用于排序。
    import re

    words: set[str] = set()
    for run in re.findall(r"[\u4e00-\u9fff]+", question):
        for n in (2, 3):
            for i in range(len(run) - n + 1):
                gram = run[i : i + n]
                if any(sw in gram for sw in _KB_STOPWORDS):
                    continue
                words.add(gram)
    words |= {w for w in re.findall(r"[A-Za-z0-9]{2,}", question)}
    if not words:
        return []
    scored: list[tuple[int, Path, str, str]] = []
    for md in sorted(kb_dir.glob("*.md")):
        try:
            text = md.read_text(encoding="utf-8")
        except OSError:
            continue
        score = sum(text.count(w) for w in words)
        if score == 0:
            continue
        # 标题：首个 H1；片段：首个命中行的上下文
        title = md.stem
        for line in text.splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
                break
        snippet = ""
        lines = text.splitlines()
        for i, line in enumerate(lines):
            if any(w in line for w in words):
                snippet = "\n".join(x.strip() for x in lines[i : i + 3] if x.strip())[:220]
                break
        scored.append((score, md, title, snippet))
    scored.sort(key=lambda x: -x[0])
    return [
        {"title": title, "source": md.name, "snippet": snippet}
        for _, md, title, snippet in scored[:limit]
    ]


def _build_ask_messages(
    member: dict[str, Any] | None,
    question: str,
    materials: list[dict],
    kb_results: list[dict],
) -> list[dict[str, str]]:
    persona = (
        member.get("prompt") or f"你是{member.get('name', '业务专家')}，{member.get('role', '')}"
        if member
        else "你是银行运营业务专家。"
    )
    sys_parts = [
        persona,
        "当前场景：客户经理在办理业务过程中向你请教，请用中文简洁、专业地回答；"
        "无法确定时明确说明，不要编造。",
    ]
    if materials:
        names = "、".join(str(f.get("file_name", "")) for f in materials[:10])
        sys_parts.append(f"已向你提交的材料：{names}")
    if kb_results:
        kb_lines = ["内部知识库检索结果（回答必须优先依据这些制度内容，并注明出处）："]
        for r in kb_results:
            kb_lines.append(f"《{r['title']}》（{r['source']}）：{r['snippet']}")
        sys_parts.append("\n".join(kb_lines))
    return [
        {"role": "system", "content": "\n\n".join(sys_parts)},
        {"role": "user", "content": question},
    ]


def _sessions_store():
    """会话存储（测试可 monkeypatch，避免碰真实 %APPDATA% sessions.db）。"""
    from agent.sessions.storage import SessionStorage

    return SessionStorage()


def _archive_qa_best_effort(
    store: CaseStorage, case_id: str, member_key: str, question: str, answer: str
) -> None:
    """会话管理归档（2026-08-11）：专家团对话也进 sessions，不再「永远只有 1 个」。

    best-effort：归档失败只记日志，绝不阻塞问答主链路。
    """
    try:
        sstore = _sessions_store()
        sid = store.get_case_session(case_id)
        if sid is None or sstore.get_session(sid) is None:
            sess = sstore.create_session(
                title=f"专家问答 · {case_id}",
                metadata={"source": "ops_case", "case_id": case_id},
            )
            sid = sess.id
            store.set_case_session(case_id, sid)
        sstore.append_message(sid, "user", question, metadata={"member_key": member_key})
        sstore.append_message(sid, "assistant", answer, metadata={"member_key": member_key})
    except Exception as e:
        logger.warning("[ops-case] sessions archive failed: %s", e)


@router.post("/case/ask")
async def case_ask(req: CaseAskRequest) -> dict:
    store = _get_case_storage()
    member = _resolve_member(req.team_id, req.member_key)
    materials = [f for f in store.list_files(req.case_id) if f.get("member_key") == req.member_key]
    # 模板/清单类请求（BUGFIX #78）：不在气泡里砸一大段文字，
    # 而是生成可界面直填的结构化草稿；生成失败自动降级普通问答
    if _is_template_request(req.question):
        draft = await _try_create_draft(store, req, member, materials)
        if draft is not None:
            answer = (
                f"已为你生成《{draft['title']}》填写表单："
                "在下方「交付草稿」区直接填写，提交后我来审核验收，通过即入交付物。"
            )
            qa = store.add_qa(
                case_id=req.case_id,
                member_key=req.member_key,
                question=req.question,
                answer=answer,
            )
            _archive_qa_best_effort(store, req.case_id, req.member_key, req.question, answer)
            await _audit(
                "ops_case_ask",
                {
                    "case_id": req.case_id,
                    "member_key": req.member_key,
                    "question": req.question[:200],
                    "draft_id": draft["id"],
                },
            )
            return {"qa": qa, "draft": _normalize_draft_row(draft)}
    # 找制度变问助手（2026-08-10）：先查内部知识库，命中则回答附制度出处（防黑盒）
    kb_results = _search_local_knowledge(req.question)
    messages = _build_ask_messages(member, req.question, materials, kb_results)
    llm_call = _make_summarize_llm()
    try:
        answer = await llm_call(messages)
    except RuntimeError as e:
        raise HTTPException(502, str(e)) from e
    answer_text = answer.strip()
    if kb_results:
        sources = "\n".join(
            f"{i + 1}. 《{r['title']}》（knowledge-base/{r['source']}）"
            for i, r in enumerate(kb_results)
        )
        answer_text += f"\n\n---\n📚 制度出处：\n{sources}"
    # 清单形态长文回答 → 自动转可直填草稿（BUGFIX #82）：问题没带「模板/清单」
    # 关键词（如「需要什么材料」）也能出表单，不在气泡里砸一大段文字；
    # 转换失败保留原文不断链
    draft: dict | None = None
    if _looks_like_checklist(answer_text):
        draft = await _try_convert_answer_to_draft(store, req, materials, answer_text)
        if draft is not None:
            answer_text = (
                f"已把资料清单生成为《{draft['title']}》草稿表单："
                "在下方「交付草稿」区直接填写，提交后我来审核验收，通过即入交付物。"
            )
            if kb_results:
                # 转草稿后仍保留制度出处（防黑盒）
                sources = "\n".join(
                    f"{i + 1}. 《{r['title']}》（knowledge-base/{r['source']}）"
                    for i, r in enumerate(kb_results)
                )
                answer_text += f"\n\n---\n📚 制度出处：\n{sources}"
    qa = store.add_qa(
        case_id=req.case_id,
        member_key=req.member_key,
        question=req.question,
        answer=answer_text,
    )
    _archive_qa_best_effort(store, req.case_id, req.member_key, req.question, answer_text)
    await _audit(
        "ops_case_ask",
        {"case_id": req.case_id, "member_key": req.member_key, "question": req.question[:200]},
    )
    if draft is not None:
        return {"qa": qa, "draft": _normalize_draft_row(draft)}
    return {"qa": qa}


# ---------------------------------------------------------------------------
# 交付草稿：要「模板/清单」时不再输出一大段文字，而是生成可界面直填的
# 结构化表单；提交后自动成为材料走专家审核，通过即并入交付物（BUGFIX #78）。
# ---------------------------------------------------------------------------

_TEMPLATE_KEYWORDS = ("模板", "清单", "表格", "样表", "样例", "范本", "空白表")

_ALLOWED_FIELD_TYPES = ("text", "textarea", "select", "date", "file")


def _is_template_request(question: str) -> bool:
    return any(kw in question for kw in _TEMPLATE_KEYWORDS)


def _build_draft_template_messages(
    member: dict[str, Any] | None,
    question: str,
    materials: list[dict],
) -> list[dict[str, str]]:
    persona = (
        member.get("prompt") or f"你是{member.get('name', '业务专家')}，{member.get('role', '')}"
        if member
        else "你是银行运营业务专家。"
    )
    sys_parts = [
        persona,
        "当前场景：客户经理需要一份可在界面直接填写的交付草稿表单（模板/清单）。"
        "请只返回 JSON，不要输出 JSON 以外的任何内容：\n"
        '{"title": "表单标题", "fields": ['
        '{"name": "小写英文字段名", "label": "中文标签", "type": "text|textarea|select|date|file", '
        '"options": ["仅 type=select 时给出"], "hint": "填写提示（可空）", "required": true}]}\n'
        "要求：字段 5-15 个（优先精简到 8 个以内，只留客户必须提供的），覆盖用户请求的清单要点；需要客户提供的信息与银行内部核查信息分开设字段；"
        "需要的是证件/文件/扫描件/照片（如身份证、营业执照、合同、报表文件）时必须用 type=file；"
        "不确定具体值时不要编造，留空由用户填写。",
    ]
    if materials:
        names = "、".join(str(f.get("file_name", "")) for f in materials[:10])
        sys_parts.append(f"已向你提交的材料（可据此推断字段）：{names}")
    return [
        {"role": "system", "content": "\n\n".join(sys_parts)},
        {"role": "user", "content": question},
    ]


def _parse_draft_template(raw: str) -> tuple[str, list[dict]] | None:
    """解析 LLM 返回的表单 JSON；不合法/无有效字段返 None（调用方降级普通问答）。"""
    from agent.llm.json_discipline import extract_json

    data = extract_json(raw)
    if not isinstance(data, dict):
        return None
    title = str(data.get("title", "") or "").strip()
    raw_fields = data.get("fields")
    if not isinstance(raw_fields, list):
        return None
    fields: list[dict] = []
    seen: set[str] = set()
    for f in raw_fields:
        if not isinstance(f, dict):
            continue
        name = str(f.get("name", "") or "").strip()
        label = str(f.get("label", "") or "").strip()
        if not name or not label or name in seen:
            continue
        seen.add(name)
        ftype = str(f.get("type", "text") or "text").strip()
        if ftype not in _ALLOWED_FIELD_TYPES:
            ftype = "textarea" if ftype in ("longtext", "multiline") else "text"
        options = [str(x) for x in (f.get("options") or []) if str(x).strip()][:20]
        fields.append(
            {
                "name": name,
                "label": label,
                "type": "select" if ftype == "select" and options else ftype,
                "options": options,
                "hint": str(f.get("hint", "") or "")[:200],
                "required": bool(f.get("required", False)),
            }
        )
    if not fields:
        return None
    return (title or "交付草稿", fields[:20])


async def _try_create_draft(
    store: CaseStorage,
    req: CaseAskRequest,
    member: dict[str, Any] | None,
    materials: list[dict],
) -> dict | None:
    """模板请求 → LLM 产出结构化表单并存草稿；任何失败返 None（绝不阻塞问答主链路）。"""
    messages = _build_draft_template_messages(member, req.question, materials)
    try:
        llm_call = _make_summarize_llm()
        raw = await llm_call(messages)
    except Exception as e:
        logger.warning("[ops-case] draft template LLM unavailable: %s", e)
        return None
    parsed = _parse_draft_template(str(raw or ""))
    if parsed is None:
        logger.warning("[ops-case] draft template unparsable, fallback to plain ask")
        return None
    title, fields = parsed
    # 已上传材料的提取要素自动预填（BUGFIX #79：少打字，用户只核不改）
    values = _prefill_draft_values(materials, fields)
    try:
        return store.add_draft(
            case_id=req.case_id,
            team_id=req.team_id,
            member_key=req.member_key,
            title=title,
            template_json=dump_json(fields),
            values_json=dump_json(values),
        )
    except Exception as e:
        logger.warning("[ops-case] draft save failed: %s", e)
        return None


def _looks_like_checklist(text: str) -> bool:
    """回答是否为清单形态（≥3 条列表项且有一定长度）：草稿转换触发条件。"""
    import re

    items = re.findall(r"^\s*(?:[-*•]|\d+[.、)])\s+\S", text, flags=re.M)
    return len(items) >= 3 and len(text) >= 100


async def _try_convert_answer_to_draft(
    store: CaseStorage,
    req: CaseAskRequest,
    materials: list[dict],
    answer_text: str,
) -> dict | None:
    """长文清单回答 → 结构化表单草稿（BUGFIX #82）；任何失败返 None（保留原文）。"""
    messages = [
        {
            "role": "system",
            "content": (
                "请把下面的专家回答转换成可在界面直接填写的表单。只返回 JSON（不要代码块）："
                '{"title": "表单标题", "fields": [{"name": "小写英文字段名", "label": "中文标签", '
                '"type": "text|textarea|select|date|file", "options": ["仅 select 需要"], '
                '"hint": "填写提示（可空）", "required": true}]}。'
                "要求：把回答中需要客户提供的每项材料/信息转成一个字段，字段数控制在 8 个以内（抓重点，不要穷举）；"
                "需要客户填写的才建字段，纯说明性内容不要建；"
                "需要的是证件/文件/扫描件/照片（如身份证、营业执照、合同、报表文件）时必须用 type=file；"
                "不要编造回答里没有的内容。"
            ),
        },
        # 输入截短至 1500 字（BUGFIX #88）：转换只需清单主体，长文尾部噪声去掉可显著降低生成耗时
        {"role": "user", "content": answer_text[:1500]},
    ]
    try:
        llm_call = _make_summarize_llm()
        raw = await llm_call(messages)
    except Exception as e:
        logger.warning("[ops-case] answer→draft LLM unavailable: %s", e)
        return None
    parsed = _parse_draft_template(str(raw or ""))
    if parsed is None:
        logger.warning("[ops-case] answer→draft unparsable, keep plain answer")
        return None
    title, fields = parsed
    values = _prefill_draft_values(materials, fields)
    try:
        return store.add_draft(
            case_id=req.case_id,
            team_id=req.team_id,
            member_key=req.member_key,
            title=title,
            template_json=dump_json(fields),
            values_json=dump_json(values),
        )
    except Exception as e:
        logger.warning("[ops-case] converted draft save failed: %s", e)
        return None


def _prefill_draft_values(materials: list[dict], fields: list[dict]) -> dict[str, str]:
    """把材料已提取要素映射到草稿字段（按 name/label 命中，取高置信优先）。"""
    candidates: dict[str, tuple[float, str]] = {}
    for f in materials:
        raw_fields = f.get("extracted_fields")
        if isinstance(raw_fields, str):
            try:
                raw_fields = json.loads(raw_fields or "[]")
            except (ValueError, TypeError):
                raw_fields = []
        for ef in raw_fields or []:
            if not isinstance(ef, dict):
                continue
            fname = str(ef.get("field", "")).strip()
            fvalue = str(ef.get("value", "")).strip()
            if not fname or not fvalue:
                continue
            try:
                conf = float(ef.get("confidence", 0.5))
            except (TypeError, ValueError):
                conf = 0.5
            if conf < 0.5:  # 低置信要素不预填，避免错值带入（员工只核标红项的原则延续）
                continue
            prev = candidates.get(fname)
            if prev is None or conf > prev[0]:
                candidates[fname] = (conf, fvalue)
    values: dict[str, str] = {}
    for field in fields:
        name = str(field.get("name", ""))
        label = str(field.get("label", ""))
        for cname, (_conf, cvalue) in candidates.items():
            if cname in (name, label) or cname in label or label in cname:
                values[name] = cvalue
                break
    return values


def _normalize_draft_row(row: dict | None) -> dict:
    """存储行 → API 响应：template_json/values_json 解析为对象。"""
    if row is None:
        return {}
    out = dict(row)
    for key, default in (("template_json", "[]"), ("values_json", "{}")):
        val = out.get(key)
        try:
            out[key.replace("_json", "")] = json.loads(val if isinstance(val, str) else default)
        except (ValueError, TypeError):
            out[key.replace("_json", "")] = [] if default == "[]" else {}
    return out


def _render_draft_markdown(title: str, template: list[dict], values: dict) -> str:
    """草稿 → md 文本（提交后作为材料文件落盘，审核通过并入交付物 zip）。"""
    lines = [f"# {title}", "", "> 本文件由运营工作台交付草稿生成，经专家验收后作为交付物。", ""]
    for f in template:
        name = str(f.get("name", ""))
        label = str(f.get("label", name))
        value = str(values.get(name, "") or "").strip()
        required_mark = "（必填）" if f.get("required") else ""
        lines.append(f"- **{label}**{required_mark}：{value or '【未填写】'}")
    lines += ["", f"生成时间：{time.strftime('%Y-%m-%d %H:%M:%S')}"]
    return "\n".join(lines) + "\n"


@router.get("/case/drafts")
async def list_case_drafts(case_id: str = Query(..., min_length=1)) -> dict:
    store = _get_case_storage()
    return {"drafts": [_normalize_draft_row(d) for d in store.list_drafts(case_id)]}


@router.put("/case/drafts/{draft_id}")
async def save_case_draft(draft_id: str, req: DraftSaveRequest) -> dict:
    store = _get_case_storage()
    row = store.get_draft(draft_id)
    if row is None:
        raise HTTPException(404, f"draft {draft_id} not found")
    if str(row.get("status", "")) == "passed":
        raise HTTPException(400, "草稿已验收通过，不能再修改")
    # 只保留模板里定义过的字段，值统一转字符串（防脏数据落盘）
    try:
        template = json.loads(str(row.get("template_json", "[]")))
    except (ValueError, TypeError):
        template = []
    allowed = {str(f.get("name", "")) for f in template if isinstance(f, dict)}
    clean = {str(k): str(v) for k, v in req.values.items() if k in allowed}
    updated = store.update_draft(draft_id, values_json=dump_json(clean))
    return _normalize_draft_row(updated)


@router.post("/case/drafts/{draft_id}/submit")
async def submit_case_draft(draft_id: str) -> dict:
    """提交草稿：校验必填 → 渲染 md 落盘为材料 → 自动 AI 审核；
    通过 → 草稿 passed（交付物）；打回 → 保留 submitted 可修改后重提。"""
    import base64

    store = _get_case_storage()
    row = store.get_draft(draft_id)
    if row is None:
        raise HTTPException(404, f"draft {draft_id} not found")
    status = str(row.get("status", ""))
    if status == "passed":
        raise HTTPException(400, "草稿已验收通过，无需重复提交")
    # 重提（上次被打回）：已有材料行则覆写内容重审，否则新建
    if status not in ("draft", "submitted"):
        raise HTTPException(400, f"unknown draft status: {status}")

    try:
        template = json.loads(str(row.get("template_json", "[]")))
        values = json.loads(str(row.get("values_json", "{}")))
    except (ValueError, TypeError):
        raise HTTPException(400, "草稿数据损坏，请重新生成")

    missing = [
        str(f.get("label", f.get("name", "")))
        for f in template
        if f.get("required") and not str(values.get(str(f.get("name", "")), "") or "").strip()
    ]
    if missing:
        raise HTTPException(400, "必填项未填写：" + "、".join(missing))

    title = str(row.get("title", "") or "交付草稿")
    md_text = _render_draft_markdown(title, template, values)
    file_name = f"{title}.md"
    case_id = str(row.get("case_id", ""))
    file_id = str(row.get("file_id", ""))
    existing = store.get_file(file_id) if file_id else None
    if existing is not None:
        # 覆写原文件内容，重置为待审，再走一遍审核
        Path(str(existing.get("file_path", ""))).write_text(md_text, encoding="utf-8")
        store.update_file(
            file_id,
            status=FILE_PENDING,
            review_note="",
            reviewed_by="",
            extracted_fields="[]",
            evidence="[]",
            reject_marks="[]",
        )
    else:
        content_b64 = base64.b64encode(md_text.encode("utf-8")).decode("ascii")
        path = store.save_upload(case_id, file_name, content_b64)
        new_file = store.add_file(
            case_id=case_id,
            team_id=str(row.get("team_id", "")),
            member_key=str(row.get("member_key", "")),
            file_name=file_name,
            file_path=str(path),
        )
        file_id = str(new_file["id"])
    store.update_draft(
        draft_id,
        status="submitted",
        file_id=file_id,
        # 快照 + 版本计数（BUGFIX #79）：打回后用户可对照上次提交内容修改
        last_snapshot=md_text,
        bump_submit_count=True,
    )
    await _audit(
        "ops_case_draft_submit",
        {"case_id": case_id, "draft_id": draft_id, "file_id": file_id, "title": title},
    )

    # 立即自动审核；审核不可用不阻塞提交（材料留待手动重审）
    file_out: dict = _normalize_file_row(store.get_file(file_id))
    try:
        file_out = await _do_review_file(store, file_id)
    except Exception as e:
        logger.warning("[ops-case] draft auto-review unavailable: %s", e)
    draft = store.get_draft(draft_id) or row
    if str(file_out.get("status", "")) == FILE_PASSED:
        draft = store.update_draft(draft_id, status="passed") or draft
    return {"draft": _normalize_draft_row(draft), "file": file_out}


def _build_export_md_rows(files: list[dict]) -> dict[str, list[dict]]:
    by_member: dict[str, list[dict]] = {}
    for f in files:
        by_member.setdefault(str(f.get("member_key", "未分组专家")), []).append(f)
    return by_member


_STATUS_LABEL = {
    FILE_PASSED: "✓ 通过",
    FILE_REJECTED: "✗ 打回",
    FILE_PENDING: "○ 待确认",
    FILE_REVIEWING: "⏳ 审核中",
}


@router.post("/case/export")
async def case_export(req: CaseExportRequest) -> dict:
    store = _get_case_storage()
    files = store.list_files(req.case_id)
    qa = store.list_qa(req.case_id)
    if not files:
        raise HTTPException(400, "Case 内没有任何交付文件，无法导出")

    target = Path(req.target_path)
    if not _path_parent_exists(target):
        raise HTTPException(400, f"导出目录不存在: {target.parent}")

    by_member = _build_export_md_rows(files)
    ts_str = time.strftime("%Y-%m-%d %H:%M:%S")
    # 多文档交叉比对 + 低置信清单（防呆提效，2026-08-10）
    inconsistencies, low_confidence = _crosscheck_fields(files)
    # 业务小结：优先 LLM，失败降级为确定性汇总（导出不被 LLM 阻塞）
    summary_md = _deterministic_summary(req, files, qa)
    try:
        llm_call = _make_summarize_llm()
        raw = await llm_call(
            [
                {
                    "role": "system",
                    "content": (
                        "你是银行运营工作台 AI。请根据下面的业务验收情况，写一段不超过 300 字的"
                        "业务小结（中文，客观克制，不要编造，直接输出正文）。"
                    ),
                },
                {"role": "user", "content": summary_md},
            ]
        )
        if raw.strip():
            summary_md = raw.strip()
    except Exception as e:
        logger.warning("[ops-case] export summary LLM fallback: %s", e)

    # 外部报告模板（2026-08-10）：专家团级 docx/md 模板 → 失败返 None 降级内置结构
    tpl_result = _render_external_template(
        req, by_member, qa, summary_md, ts_str, inconsistencies, low_confidence
    )

    with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zf:
        _write_export_zip_contents(
            zf,
            req,
            by_member,
            qa,
            summary_md,
            ts_str,
            inconsistencies,
            low_confidence,
            tpl_result,
        )

    await _audit(
        "ops_case_export",
        {
            "case_id": req.case_id,
            "target_path": req.target_path,
            "file_count": len(files),
        },
    )
    return {"ok": True, "path": str(target), "file_count": len(files)}


def _write_export_zip_contents(
    zf: zipfile.ZipFile,
    req: CaseExportRequest,
    by_member: dict[str, list[dict]],
    qa: list[dict],
    summary_md: str,
    ts_str: str,
    inconsistencies: list[dict],
    low_confidence: list[dict],
    tpl_result: tuple[str, bytes] | tuple[str, str] | None = None,
) -> None:
    """同步写入 zip 内容（交付文件/检查结果/问答记录/报告初稿/小结/清单）。"""
    # 1) 交付文件（按专家分目录）
    for member_key, member_files in by_member.items():
        for f in member_files:
            src = Path(str(f.get("file_path", "")))
            if src.is_file():
                zf.write(src, f"交付文件/{member_key}/{src.name}")
    # 2) 检查结果（每位专家一份 md）
    for member_key, member_files in by_member.items():
        lines = [f"# {member_key} 检查结果\n"]
        for f in member_files:
            label = _STATUS_LABEL.get(str(f.get("status", "")), str(f.get("status", "")))
            lines.append(f"## {f.get('file_name', '')} — {label}")
            note = str(f.get("review_note", "")).strip()
            lines.append(note if note else "（无审核意见）")
            lines.append("")
        zf.writestr(f"检查结果/{member_key}.md", "\n".join(lines))
    # 3) 问答记录
    if qa:
        qa_lines = ["# 专家问答记录\n"]
        for item in qa:
            qa_lines.append(f"## 问（→ {item.get('member_key', '')}）")
            qa_lines.append(str(item.get("question", "")))
            qa_lines.append("\n**答：**\n" + str(item.get("answer", "")))
            qa_lines.append("")
        zf.writestr("问答记录.md", "\n".join(qa_lines))
    # 4) 报告初稿（写八股文变改填空题，2026-08-10）：外部模板 → 内置 docx → 内置 md
    if tpl_result is not None:
        tpl_kind, tpl_content = tpl_result
        if tpl_kind == "docx":
            zf.writestr("尽调报告初稿.docx", tpl_content)
        else:
            zf.writestr("报告初稿.md", tpl_content)
    else:
        report_docx = _build_report_docx(
            req, by_member, qa, summary_md, ts_str, inconsistencies, low_confidence
        )
        if report_docx is not None:
            zf.writestr("尽调报告初稿.docx", report_docx)
        else:
            zf.writestr(
                "报告初稿.md",
                _build_report_markdown(
                    req, by_member, qa, summary_md, ts_str, inconsistencies, low_confidence
                ),
            )
    # 5) 业务小结 + 交付物清单
    zf.writestr("业务小结.md", f"# 业务小结\n\n{summary_md}\n")
    file_count = sum(len(v) for v in by_member.values())
    readme_lines = [
        "# 交付物清单\n",
        f"- 业务：{req.feature_name or req.case_id}",
        f"- 专家团：{req.team_name or '未指定'}",
        f"- 导出时间：{ts_str}",
        f"- 材料文件：{file_count} 份",
        "",
        "## 验收清单",
    ]
    if req.checklist:
        readme_lines += [f"- [x] {x}" for x in req.checklist]
    else:
        readme_lines.append("- （无）")
    zf.writestr("README.md", "\n".join(readme_lines))


def _deterministic_summary(req: CaseExportRequest, files: list[dict], qa: list[dict]) -> str:
    lines = [f"业务：{req.feature_name or req.case_id}"]
    passed = sum(1 for f in files if f.get("status") == FILE_PASSED)
    lines.append(f"材料共 {len(files)} 份，验收通过 {passed} 份。")
    for f in files:
        label = _STATUS_LABEL.get(str(f.get("status", "")), str(f.get("status", "")))
        lines.append(f"- {f.get('member_key', '')} / {f.get('file_name', '')}：{label}")
    if req.checklist:
        lines.append("验收清单：" + "；".join(req.checklist))
    if qa:
        lines.append(f"专家问答 {len(qa)} 条。")
    return "\n".join(lines)


def _build_template_context(
    req: CaseExportRequest,
    by_member: dict[str, list[dict]],
    qa: list[dict],
    summary_md: str,
    ts_str: str,
    inconsistencies: list[dict],
    low_confidence: list[dict],
):
    """组装外部报告模板的渲染上下文（占位符规范见 ops/report_template.py）。"""
    from .report_template import TemplateContext

    files = [f for group in by_member.values() for f in group]

    acceptance_lines: list[str] = []
    rows: list[dict[str, str]] = []
    for member_key, group in by_member.items():
        for f in group:
            label = _STATUS_LABEL.get(str(f.get("status", "")), str(f.get("status", "")))
            note = str(f.get("review_note", "")).strip()
            line = f"- {member_key} / {f.get('file_name', '')}：{label}"
            if note:
                line += f"（意见：{note[:120]}）"
            acceptance_lines.append(line)
            rows.append(
                {
                    "材料": str(f.get("file_name", "")),
                    "专家": member_key,
                    "状态": label,
                    "意见": note[:200],
                }
            )

    cross_lines: list[str] = []
    if inconsistencies:
        for inc in inconsistencies:
            detail = "；".join(f"{v['file']}={v['value']}" for v in inc["values"])
            cross_lines.append(f"- ⚠ 字段「{inc['field']}」不一致：{detail}")
    else:
        cross_lines.append("- 未发现跨材料要素不一致")

    qa_lines = [
        f"- 问（→ {item.get('member_key', '')}）：{item.get('question', '')} "
        f"答：{str(item.get('answer', ''))[:200]}"
        for item in qa
    ] or ["- （无问答记录）"]

    manual_lines = [f"- {x}" for x in _report_manual_items(files, low_confidence)]

    return TemplateContext(
        text={
            "业务名称": req.feature_name or req.case_id,
            "专家团": req.team_name or "未指定",
            "生成时间": ts_str,
            "材料数量": str(len(files)),
            "风险结论": summary_md,
        },
        lists={
            "材料验收清单": acceptance_lines,
            "交叉比对清单": cross_lines,
            "问答记录清单": qa_lines,
            "人工确认事项": manual_lines,
        },
        rows=rows,
    )


def _render_external_template(
    req: CaseExportRequest,
    by_member: dict[str, list[dict]],
    qa: list[dict],
    summary_md: str,
    ts_str: str,
    inconsistencies: list[dict],
    low_confidence: list[dict],
) -> tuple[str, bytes] | tuple[str, str] | None:
    """外部报告模板渲染：('docx', bytes) / ('md', str) / None（降级内置结构）。

    红线：模板是可选增强，任何异常都返 None，绝不阻塞导出。
    """
    if not req.team_id:
        return None
    try:
        from agent.expert_teams.api import get_loader
        from agent.expert_teams.templates import resolve_template_path

        team = get_loader().get(req.team_id)
        if team is None:
            return None
        tpl_path = resolve_template_path(team.id, team.report_template)
        if tpl_path is None:
            return None

        from .report_template import render_docx_template, render_markdown_template

        ctx = _build_template_context(
            req, by_member, qa, summary_md, ts_str, inconsistencies, low_confidence
        )
        if tpl_path.suffix.lower() == ".docx":
            data = render_docx_template(tpl_path, ctx)
            if data:
                return ("docx", data)
            return None
        return ("md", render_markdown_template(tpl_path, ctx))
    except Exception as e:
        logger.warning("[ops-case] external template render failed, fallback builtin: %s", e)
        return None


# ---------------------------------------------------------------------------
# 多文档交叉比对 + 报告初稿（2026-08-10）
# ---------------------------------------------------------------------------


def _loads_json_list(value: Any) -> list[dict]:
    """存储行的 JSON 文本 → list（兼容已解析的 list）。"""
    if isinstance(value, list):
        return [x for x in value if isinstance(x, dict)]
    if isinstance(value, str):
        try:
            data = json.loads(value or "[]")
            if isinstance(data, list):
                return [x for x in data if isinstance(x, dict)]
        except (ValueError, TypeError):
            pass
    return []


def _crosscheck_fields(files: list[dict]) -> tuple[list[dict], list[dict]]:
    """多文档交叉比对（人肉比对变自动交叉，2026-08-10）。

    返回 (inconsistencies, low_confidence)：
    - inconsistencies：同名字段在不同材料中取值不一致（如法人名字对不上）
    - low_confidence：置信度 < 0.6 的要素清单（需人工核对）
    """
    by_field: dict[str, list[dict]] = {}
    low: list[dict] = []
    for f in files:
        fname = str(f.get("file_name", ""))
        for item in _loads_json_list(f.get("extracted_fields")):
            name = str(item.get("field", "")).strip()
            value = str(item.get("value", "")).strip()
            if not name or not value:
                continue
            try:
                conf = float(item.get("confidence", 0.8))
            except (TypeError, ValueError):
                conf = 0.8
            by_field.setdefault(name, []).append(
                {"file": fname, "value": value, "confidence": conf}
            )
            if conf < 0.6:
                low.append({"file": fname, "field": name, "value": value, "confidence": conf})
    inconsistencies: list[dict] = []
    for name, occs in by_field.items():
        if len(occs) >= 2 and len({o["value"] for o in occs}) > 1:
            inconsistencies.append({"field": name, "values": occs})
    return inconsistencies, low


@router.get("/case/crosscheck")
async def case_crosscheck(case_id: str = Query(...)) -> dict:
    """同 Case 内多份材料的要素交叉比对（不一致标红，防退件）。"""
    store = _get_case_storage()
    files = store.list_files(case_id)
    inconsistencies, low_confidence = _crosscheck_fields(files)
    return {
        "case_id": case_id,
        "inconsistencies": inconsistencies,
        "low_confidence": low_confidence,
        "consistent": len(inconsistencies) == 0,
    }


def _report_manual_items(files: list[dict], low_confidence: list[dict]) -> list[str]:
    """报告「需人工确认」清单（人机边界：AI 干脏活，人干细活）。"""
    items: list[str] = []
    for f in files:
        status = str(f.get("status", ""))
        if status in (FILE_REJECTED, FILE_PENDING):
            label = _STATUS_LABEL.get(status, status)
            items.append(f"{f.get('file_name', '')}：验收状态为{label}，需人工复核后改判")
    for lc in low_confidence:
        items.append(
            f"{lc['file']}：要素「{lc['field']}」置信度仅 "
            f"{lc['confidence'] * 100:.0f}%（值：{lc['value']}），需人工核对"
        )
    items.append("报告中的主观判断与风险定性段落（标注【需人工填写】处）")
    items.append("签字确认：经办人 / 复核人 / 日期")
    return items


def _build_report_markdown(
    req: CaseExportRequest,
    by_member: dict[str, list[dict]],
    qa: list[dict],
    summary_md: str,
    ts_str: str,
    inconsistencies: list[dict],
    low_confidence: list[dict],
) -> str:
    """报告初稿的 md 降级版（python-docx 不可用时）。"""
    files = [f for group in by_member.values() for f in group]
    lines = [
        f"# {req.feature_name or req.case_id} 尽调报告（初稿）\n",
        "> 本报告为 AI 生成初稿：客观数据已自动填充；标注【需人工填写】处必须由人工确认补全。\n",
        "## 一、业务概况",
        f"- 业务：{req.feature_name or req.case_id}",
        f"- 专家团：{req.team_name or '未指定'}",
        f"- 生成时间：{ts_str}",
        f"- 材料：{len(files)} 份\n",
        "## 二、材料验收结果",
    ]
    for member_key, group in by_member.items():
        for f in group:
            label = _STATUS_LABEL.get(str(f.get("status", "")), str(f.get("status", "")))
            lines.append(f"- {member_key} / {f.get('file_name', '')}：{label}")
            note = str(f.get("review_note", "")).strip()
            if note:
                lines.append(f"  - 意见：{note}")
    lines += ["", "## 三、交叉比对结果"]
    if inconsistencies:
        for inc in inconsistencies:
            detail = "；".join(f"{v['file']}={v['value']}" for v in inc["values"])
            lines.append(f"- ⚠ 字段「{inc['field']}」不一致：{detail}")
    else:
        lines.append("- 未发现跨材料要素不一致")
    lines += ["", "## 四、风险与结论（AI 初稿，需人工审核）", summary_md, ""]
    lines += ["## 五、需人工确认事项"]
    lines += [f"- {x}" for x in _report_manual_items(files, low_confidence)]
    lines += ["", "经办人：【需人工填写】　复核人：【需人工填写】　日期：【需人工填写】"]
    return "\n".join(lines) + "\n"


def _build_report_docx(
    req: CaseExportRequest,
    by_member: dict[str, list[dict]],
    qa: list[dict],
    summary_md: str,
    ts_str: str,
    inconsistencies: list[dict],
    low_confidence: list[dict],
) -> bytes | None:
    """尽调报告初稿 docx（写八股文变改填空题，2026-08-10）。

    python-docx 不可用/生成失败 → 返 None，由调用方降级 md。
    人机边界：客观数据自动填充；低置信/待确认项标红；主观判断留「【需人工填写】」空位。
    """
    try:
        import io

        from docx import Document
        from docx.shared import RGBColor
    except Exception as e:
        logger.warning("[ops-case] python-docx unavailable, report fallback to md: %s", e)
        return None

    red = RGBColor(0xCD, 0x31, 0x31)
    files = [f for group in by_member.values() for f in group]
    try:
        doc = Document()
        doc.add_heading(f"{req.feature_name or req.case_id} 尽调报告（初稿）", level=0)
        meta = doc.add_paragraph()
        meta.add_run(
            "本报告为 AI 生成初稿：客观数据已自动填充；"
            "标注【需人工填写】处必须由人工确认补全，签字后方可提交。"
        ).italic = True

        doc.add_heading("一、业务概况", level=1)
        for line in (
            f"业务：{req.feature_name or req.case_id}",
            f"专家团：{req.team_name or '未指定'}",
            f"生成时间：{ts_str}",
            f"材料：{len(files)} 份",
        ):
            doc.add_paragraph(line)

        doc.add_heading("二、材料验收结果", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, head in enumerate(("材料", "审核专家", "状态", "AI 意见")):
            table.rows[0].cells[i].text = head
        for member_key, group in by_member.items():
            for f in group:
                row = table.add_row().cells
                row[0].text = str(f.get("file_name", ""))
                row[1].text = member_key
                row[2].text = _STATUS_LABEL.get(str(f.get("status", "")), str(f.get("status", "")))
                row[3].text = str(f.get("review_note", ""))[:200]

        doc.add_heading("三、交叉比对结果", level=1)
        if inconsistencies:
            for inc in inconsistencies:
                p = doc.add_paragraph()
                run = p.add_run(
                    "⚠ 字段「"
                    + str(inc["field"])
                    + "」不一致："
                    + "；".join(f"{v['file']}={v['value']}" for v in inc["values"])
                )
                run.font.color.rgb = red
        else:
            doc.add_paragraph("未发现跨材料要素不一致。")

        doc.add_heading("四、专家问答记录", level=1)
        if qa:
            for item in qa:
                doc.add_paragraph(
                    f"问（→ {item.get('member_key', '')}）：{item.get('question', '')}"
                )
                doc.add_paragraph(f"答：{str(item.get('answer', ''))[:500]}")
        else:
            doc.add_paragraph("（无问答记录）")

        doc.add_heading("五、风险与结论（AI 初稿，需人工审核）", level=1)
        doc.add_paragraph(summary_md)

        doc.add_heading("六、需人工确认事项", level=1)
        for item in _report_manual_items(files, low_confidence):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item)

        doc.add_paragraph("经办人：【需人工填写】　复核人：【需人工填写】　日期：【需人工填写】")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        logger.warning("[ops-case] report docx build failed, fallback to md: %s", e)
        return None
