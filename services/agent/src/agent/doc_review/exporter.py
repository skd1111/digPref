"""审核结果导出 Word（.docx）。

两种模式：
  - risks_only：结构化风险报告（风险点 + 对应原文摘录 + 知识库依据）
  - full：导出文档全部原文，风险点以 Word 批注（comment）形式挂在对应原文上

批注实现：python-docx 无原生 comment API，此处直接构造 OOXML：
  word/comments.xml 部件 + 正文 commentRangeStart/End + commentReference。
"""

from __future__ import annotations

import io
import json
from datetime import datetime
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree

_AUTHOR = "EAIDE 审核专家"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_LEVEL_CN = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危"}
_LEVEL_COLOR = {
    "critical": RGBColor(0xB3, 0x26, 0x1E),
    "high": RGBColor(0xCD, 0x31, 0x31),
    "medium": RGBColor(0xB2, 0x5C, 0x1A),
    "low": RGBColor(0x05, 0x96, 0x69),
}
_LEVEL_ORDER = {"critical": 3, "high": 2, "medium": 1, "low": 0}


def _sorted_findings(findings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(
        findings,
        key=lambda f: _LEVEL_ORDER.get(str(f.get("risk_level", "")), -1),
        reverse=True,
    )


def _level_cn(level: str) -> str:
    return _LEVEL_CN.get(level, level)


def _parse_positions(finding: dict[str, Any]) -> list[dict[str, Any]]:
    raw = finding.get("positions_json") or finding.get("positions") or []
    if isinstance(raw, str):
        try:
            raw = json.loads(raw)
        except json.JSONDecodeError:
            return []
    return [p for p in raw if isinstance(p, dict)]


# ---- 批注（full 模式）-------------------------------------------------------


def _add_comment_content(comments_el: etree._Element, cid: int, text: str) -> None:
    """向 comments 根元素追加一条 <w:comment>（多段落按换行拆分）。"""
    c = etree.SubElement(comments_el, qn("w:comment"))
    c.set(qn("w:id"), str(cid))
    c.set(qn("w:author"), _AUTHOR)
    c.set(qn("w:date"), datetime.now().isoformat() + "Z")
    for line in text.split("\n"):
        p = etree.SubElement(c, qn("w:p"))
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = line


def _attach_comments_part(doc: Document, comments_el: etree._Element) -> None:
    """把 comments XML 挂为 word/comments.xml 部件并建立关系。"""
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    xml_bytes = etree.tostring(comments_el, xml_declaration=True, encoding="UTF-8", standalone=True)
    part = Part(
        PackURI("/word/comments.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
        xml_bytes,
        doc.part.package,
    )
    doc.part.relate_to(
        part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    )


def _comment_body(finding: dict[str, Any]) -> str:
    """批注正文：标题 + 等级 + 描述（+ 建议），控制在可读长度。"""
    lines = [
        f"【{_level_cn(str(finding.get('risk_level', '')))} · "
        f"{finding.get('risk_type', '')}】{finding.get('title', '')}"
    ]
    desc = str(finding.get("description", "") or "").strip()
    if desc:
        lines.append(desc)
    sug = str(finding.get("suggestion", "") or "").strip()
    if sug:
        lines.append(f"建议：{sug}")
    return "\n".join(lines)


def _emit_block_paragraph(
    doc: Document,
    text: str,
    intervals: list[tuple[int, int, int]],
) -> None:
    """渲染一个文本块：原文照排，批注区间用 commentRangeStart/End 包裹。

    intervals: (start, end, comment_id)，允许重叠。
    """
    para = doc.add_paragraph()
    p_el = para._p
    if not intervals or not text:
        run = para.add_run(text)
        run.font.size = Pt(11)
        return

    points = {0, len(text)}
    for s, e, _cid in intervals:
        points.add(max(0, min(s, len(text))))
        points.add(max(0, min(e, len(text))))
    pts = sorted(points)
    for i in range(len(pts) - 1):
        a, b = pts[i], pts[i + 1]
        for s, _e, cid in intervals:
            if s == a:
                el = p_el.makeelement(qn("w:commentRangeStart"), {})
                el.set(qn("w:id"), str(cid))
                p_el.append(el)
        if b > a:
            r = para.add_run(text[a:b])
            r.font.size = Pt(11)
        for _s, e, cid in intervals:
            if e == b:
                el = p_el.makeelement(qn("w:commentRangeEnd"), {})
                el.set(qn("w:id"), str(cid))
                p_el.append(el)
                ref_run = p_el.makeelement(qn("w:r"), {})
                ref = p_el.makeelement(qn("w:commentReference"), {})
                ref.set(qn("w:id"), str(cid))
                ref_run.append(ref)
                p_el.append(ref_run)


def _build_full(parsed: dict[str, Any], findings: list[dict[str, Any]]) -> Document:
    """全文模式：按页/块重建原文，finding 依据 positions 挂批注。"""
    doc = Document()
    doc.add_heading(f"{parsed.get('file_name', '文档')} — 风险审核（全文批注版）", level=0)
    meta = doc.add_paragraph()
    meta.add_run(
        f"共 {len(findings)} 个风险点，以 Word 批注形式标注在对应原文处；"
        f"审阅请在 Word 中打开批注窗格。"
    ).italic = True

    # block_id → 块元信息（含全文偏移 start，用于坐标换算）
    block_index: dict[str, dict[str, Any]] = {
        str(block.get("block_id", "")): block
        for page in parsed.get("pages", [])
        for block in page.get("blocks", [])
    }

    # block_id → [(start, end, comment_id)]，坐标为块内局部偏移
    # 注意：finding.positions 的 start/end 是全文偏移（matcher.py 语义），
    # 必须减去 block.start 换算成块内坐标；跨块证据在每个覆盖块内各挂一段
    comments_el = etree.Element(qn("w:comments"), nsmap={"w": _W_NS})
    block_comments: dict[str, list[tuple[int, int, int]]] = {}
    cid = 0
    for finding in _sorted_findings(findings):
        positions = _parse_positions(finding)
        if not positions:
            continue  # 无法定位原文的 finding 不挂批注（避免噪声）
        _add_comment_content(comments_el, cid, _comment_body(finding))
        for pos in positions:
            bid = str(pos.get("block_id", ""))
            g_start = int(pos.get("start", 0))
            g_end = int(pos.get("end", 0))
            if g_end <= g_start or not bid:
                continue
            block = block_index.get(bid)
            if block is None:
                continue
            # 全文偏移 → 块内局部偏移（跨块时裁剪到本块范围）
            local_start = max(0, g_start - block["start"])
            local_end = min(len(block["text"]), g_end - block["start"])
            if local_end > local_start:
                block_comments.setdefault(bid, []).append((local_start, local_end, cid))
        cid += 1

    for page in parsed.get("pages", []):
        pno = page.get("page_no", "")
        doc.add_heading(f"第 {pno} 页", level=1)
        for block in page.get("blocks", []):
            text = str(block.get("text", ""))
            if not text.strip():
                continue
            _emit_block_paragraph(doc, text, block_comments.get(str(block.get("block_id", "")), []))

    if cid > 0:
        _attach_comments_part(doc, comments_el)
    return doc


# ---- 结构化风险报告（risks_only 模式）---------------------------------------


def _build_risks_only(
    parsed: dict[str, Any], run: dict[str, Any], findings: list[dict[str, Any]]
) -> Document:
    doc = Document()
    doc.add_heading(f"{parsed.get('file_name', '文档')} — 风险审核报告", level=0)

    sorted_f = _sorted_findings(findings)
    counts = {lv: sum(1 for f in sorted_f if f.get("risk_level") == lv) for lv in _LEVEL_CN}

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    info = [
        ("文件名", str(parsed.get("file_name", ""))),
        ("文档类型", str(run.get("doc_category", "") or "—")),
        ("总体风险等级", _level_cn(str(run.get("overall_risk_level", "") or ""))),
        (
            "风险点统计",
            f"共 {len(sorted_f)} 条：严重 {counts['critical']} · 高危 {counts['high']} · "
            f"中危 {counts['medium']} · 低危 {counts['low']}",
        ),
    ]
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    if run.get("summary"):
        doc.add_paragraph(str(run["summary"]))

    if not sorted_f:
        doc.add_paragraph("未发现风险点。")
        return doc

    doc.add_heading("风险明细", level=1)
    for idx, f in enumerate(sorted_f, start=1):
        level = str(f.get("risk_level", ""))
        h = doc.add_heading(f"风险 {idx} · {f.get('title', '')}（{_level_cn(level)}）", level=2)
        if h.runs:
            h.runs[0].font.color.rgb = _LEVEL_COLOR.get(level, RGBColor(0x1F, 0x1F, 0x1F))

        tag = doc.add_paragraph()
        tag_run = tag.add_run(f"{f.get('risk_type', '')} · {level}")
        tag_run.font.color.rgb = _LEVEL_COLOR.get(level, RGBColor(0x61, 0x61, 0x61))
        tag_run.font.bold = True

        desc = str(f.get("description", "") or "").strip()
        if desc:
            doc.add_paragraph().add_run("风险描述：").bold = True
            doc.add_paragraph(desc)
        sug = str(f.get("suggestion", "") or "").strip()
        if sug:
            doc.add_paragraph().add_run("处理建议：").bold = True
            doc.add_paragraph(sug)
        evidence = str(f.get("evidence_text", "") or "").strip()
        if evidence:
            doc.add_paragraph().add_run("原文摘录：").bold = True
            q = doc.add_paragraph(evidence)
            q.paragraph_format.left_indent = Pt(18)
            if q.runs:
                q.runs[0].font.color.rgb = RGBColor(0x61, 0x61, 0x61)
        if f.get("rule_ref"):
            doc.add_paragraph(f"规则依据：{f['rule_ref']}")
        kb_refs = f.get("kb_refs") or []
        if kb_refs:
            doc.add_paragraph().add_run("知识库依据：").bold = True
            for ref in kb_refs:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{ref.get('source', '')} · {ref.get('heading', '')}").bold = True
                excerpt = str(ref.get("excerpt", "") or "").strip()
                if excerpt:
                    doc.add_paragraph(excerpt).paragraph_format.left_indent = Pt(36)
    return doc


# ---- 入口 -------------------------------------------------------------------


def build_export_docx(
    *,
    mode: str,
    parsed: dict[str, Any],
    findings: list[dict[str, Any]],
    run: dict[str, Any] | None = None,
) -> bytes:
    """生成 docx 二进制。mode = 'full'（全文+批注）| 'risks_only'（结构化报告）。"""
    if mode == "full":
        doc = _build_full(parsed, findings)
    elif mode == "risks_only":
        doc = _build_risks_only(parsed, run or {}, findings)
    else:
        raise ValueError(f"unsupported export mode: {mode}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
