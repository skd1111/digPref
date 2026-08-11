"""Phase 1B V6 · 文档处理工具族（document_tool suite，2026-08-10）。

对照 Work/Coding Agent 通用内置工具清单，补齐「文档处理」类的结构化读写能力
（V5 file_to_markdown 只覆盖"读"，本模块覆盖"查 / 聚合 / 生成 / 合并拆分"）：

工具清单：
    - excel_query    Excel / CSV 结构化查询（sheets / rows / aggregate；只读，read）
    - excel_export   JSON 行数据 → xlsx（写文件，medium → HITL）
    - pdf_merge      多个 PDF 合并为一个（写文件，medium → HITL）
    - pdf_split      PDF 按页码区间拆分（写文件，medium → HITL）
    - word_generate  Markdown 子集文本 → Word 文档（写文件，medium → HITL）

安全约束（与 V0-V5 一致）：
    - 所有路径先走 path_sandbox.validate_path()（Windows 保留名 / UNC / null byte）。
    - 写工具输出路径默认禁止覆盖（overwrite=True 显式放行）。
    - openpyxl 是可选依赖（data-full extra）：缺失时返友好错误而不是崩溃。
    - pypdf / python-docx 是 agent 主依赖，直接延迟导入。
"""

from __future__ import annotations

import csv as _csv
import re
from pathlib import Path
from typing import Any

from agent.builtin.models import ToolResult
from agent.builtin.path_sandbox import validate_path

# 查询 / 输出上限（防大文件爆内存）
_MAX_QUERY_ROWS = 500
_MAX_EXPORT_ROWS = 10_000
_MAX_CELL_CHARS = 100_000
_MAX_MERGE_INPUTS = 50

_AGG_OPS = frozenset({"count", "sum", "avg", "min", "max"})


def _missing_dep_result(dep: str, extra: str, risk_level: str) -> ToolResult:
    return ToolResult(
        ok=False,
        error="missing_dependency",
        hint=f"缺少可选依赖 {dep}，请安装 agent[{extra}]（内网可用 config/driver 离线 wheel）",
        risk_level=risk_level,  # type: ignore[arg-type]
    )


def _cell_value(v: Any) -> Any:
    """单元格 / 值 → JSON 安全值（Cell 对象取 .value；datetime / date → ISO 字符串）。"""
    if hasattr(v, "value") and hasattr(v, "column"):  # openpyxl Cell / ReadOnlyCell
        v = v.value
    if hasattr(v, "isoformat"):
        return v.isoformat()
    return v


def _read_csv_rows(path: Path) -> tuple[list[str], list[list[Any]]]:
    """CSV → (header, rows)。空文件返 ([], [])。"""
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as fh:
        reader = _csv.reader(fh)
        all_rows = [[_cell_value(c) for c in row] for row in reader]
    if not all_rows:
        return [], []
    return [str(c) for c in all_rows[0]], all_rows[1:]


def _load_sheet(path: Path, sheet: str | None) -> tuple[str, list[str], list[list[Any]]]:
    """加载 xlsx 单个 sheet（或 csv 全文件）→ (sheet_name, header, rows)。"""
    if path.suffix.lower() == ".csv":
        header, rows = _read_csv_rows(path)
        return path.stem, header, rows
    from openpyxl import load_workbook  # 可选依赖

    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        name = sheet if sheet and sheet in wb.sheetnames else wb.sheetnames[0]
        ws = wb[name]
        grid = [[_cell_value(c) for c in row] for row in ws.iter_rows()]
    finally:
        wb.close()
    if not grid:
        return name, [], []
    header = [str(c) if c is not None else "" for c in grid[0]]
    return name, header, grid[1:]


def _column_index(header: list[str], column: str) -> int | None:
    """列名 → 下标（大小写不敏感）；找不到返 None。"""
    lower = column.strip().lower()
    for i, h in enumerate(header):
        if h.strip().lower() == lower:
            return i
    return None


def _numeric(v: Any) -> float | None:
    if isinstance(v, bool):
        return None
    if isinstance(v, int | float):
        return float(v)
    if isinstance(v, str):
        try:
            return float(v.replace(",", "").strip())
        except ValueError:
            return None
    return None


def builtin_excel_query(
    *,
    path: str,
    action: str = "rows",
    sheet: str | None = None,
    columns: list[str] | None = None,
    where: dict[str, Any] | None = None,
    group_by: str | None = None,
    agg_column: str | None = None,
    agg_op: str = "sum",
    limit: int = 100,
    offset: int = 0,
) -> ToolResult:
    """Excel / CSV 结构化查询（只读）。

    action:
        sheets    列出全部 sheet 名（xlsx）
        rows      取行（支持 columns 投影 / where 等值过滤 / limit+offset 分页）
        aggregate 按 group_by 列分组聚合（agg_op: count/sum/avg/min/max）
    where 形如 {"列名": 值}，多条件为 AND（字符串等值，忽略大小写）。
    """
    if action not in ("sheets", "rows", "aggregate"):
        return ToolResult(
            ok=False,
            error="invalid_action",
            hint="action 必须是 sheets / rows / aggregate",
            risk_level="read",
        )
    try:
        p = validate_path(path, must_exist=True)
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="read")
    if not p.is_file():
        return ToolResult(ok=False, error="not_a_file", hint=str(p), risk_level="read")

    try:
        if action == "sheets":
            if p.suffix.lower() == ".csv":
                return ToolResult(ok=True, content={"sheets": [p.stem]}, risk_level="read")
            from openpyxl import load_workbook

            wb = load_workbook(p, read_only=True)
            try:
                names = list(wb.sheetnames)
            finally:
                wb.close()
            return ToolResult(ok=True, content={"sheets": names}, risk_level="read")

        sheet_name, header, rows = _load_sheet(p, sheet)
    except ImportError:
        return _missing_dep_result("openpyxl", "data-full", "read")
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="read")

    if not header:
        return ToolResult(
            ok=True,
            content={"sheet": sheet_name, "header": [], "rows": [], "total": 0},
            risk_level="read",
        )

    # where 等值过滤
    if where:
        filters: list[tuple[int, str]] = []
        for col, val in where.items():
            idx = _column_index(header, col)
            if idx is None:
                return ToolResult(
                    ok=False,
                    error="unknown_column",
                    hint=f"where 列 {col!r} 不存在，表头: {header}",
                    risk_level="read",
                )
            filters.append((idx, str(val).strip().lower()))
        rows = [
            r
            for r in rows
            if all(str(r[i] if i < len(r) else "").strip().lower() == v for i, v in filters)
        ]

    if action == "aggregate":
        if not group_by:
            return ToolResult(
                ok=False,
                error="missing_group_by",
                hint="aggregate 需要 group_by 列名",
                risk_level="read",
            )
        if agg_op not in _AGG_OPS:
            return ToolResult(
                ok=False,
                error="invalid_agg_op",
                hint=f"agg_op 必须是 {sorted(_AGG_OPS)}",
                risk_level="read",
            )
        g_idx = _column_index(header, group_by)
        if g_idx is None:
            return ToolResult(
                ok=False,
                error="unknown_column",
                hint=f"group_by 列 {group_by!r} 不存在，表头: {header}",
                risk_level="read",
            )
        a_idx = _column_index(header, agg_column) if agg_column else None
        if agg_column and a_idx is None:
            return ToolResult(
                ok=False,
                error="unknown_column",
                hint=f"agg_column {agg_column!r} 不存在，表头: {header}",
                risk_level="read",
            )
        groups: dict[str, list[float]] = {}
        for r in rows:
            key = str(r[g_idx] if g_idx < len(r) else "")
            vals = groups.setdefault(key, [])
            if a_idx is not None:
                n = _numeric(r[a_idx] if a_idx < len(r) else None)
                if n is not None:
                    vals.append(n)
            else:
                vals.append(0.0)  # count 无聚合列：每行计 1
        out: list[dict[str, Any]] = []
        for key, vals in groups.items():
            entry: dict[str, Any] = {group_by: key, "count": len(vals)}
            if agg_op == "sum":
                entry["value"] = round(sum(vals), 6)
            elif agg_op == "avg":
                entry["value"] = round(sum(vals) / len(vals), 6) if vals else None
            elif agg_op == "min":
                entry["value"] = min(vals) if vals else None
            elif agg_op == "max":
                entry["value"] = max(vals) if vals else None
            out.append(entry)
        out.sort(key=lambda e: str(e[group_by]))
        return ToolResult(
            ok=True,
            content={"sheet": sheet_name, "agg_op": agg_op, "groups": out},
            meta={"group_count": len(out), "row_count": len(rows)},
            risk_level="read",
        )

    # action == "rows"
    total = len(rows)
    rows = rows[max(0, int(offset)) : max(0, int(offset)) + min(int(limit), _MAX_QUERY_ROWS)]
    if columns:
        idxs = []
        for col in columns:
            idx = _column_index(header, col)
            if idx is None:
                return ToolResult(
                    ok=False,
                    error="unknown_column",
                    hint=f"columns 列 {col!r} 不存在，表头: {header}",
                    risk_level="read",
                )
            idxs.append(idx)
        header = [header[i] for i in idxs]
        rows = [[r[i] if i < len(r) else None for i in idxs] for r in rows]
    return ToolResult(
        ok=True,
        content={"sheet": sheet_name, "header": header, "rows": rows, "total": total},
        meta={"row_count": len(rows), "total": total, "truncated": total > len(rows)},
        risk_level="read",
    )


def _check_output_path(path: str, *, overwrite: bool, risk_level: str) -> Path | ToolResult:
    """写工具输出路径校验：沙箱 + 存在性冲突检查。"""
    try:
        out = validate_path(path, must_exist=False)
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level=risk_level)  # type: ignore[arg-type]
    if out.exists() and not overwrite:
        return ToolResult(
            ok=False,
            error="exists_no_overwrite",
            hint=f"目标文件已存在: {out}（传 overwrite=True 覆盖）",
            risk_level=risk_level,  # type: ignore[arg-type]
        )
    if out.parent and not out.parent.exists():
        out.parent.mkdir(parents=True, exist_ok=True)
    return out


def builtin_excel_export(
    *,
    path: str,
    rows: list[Any],
    sheet_name: str = "Sheet1",
    overwrite: bool = False,
) -> ToolResult:
    """把 JSON 行数据写为 xlsx（medium 风险，走 HITL）。

    rows 支持两种形态：
      - list[dict]  → 表头 = 各行键的并集（按首行顺序优先）
      - list[list]  → 首行视为表头直接写
    """
    if not isinstance(rows, list) or not rows:
        return ToolResult(ok=False, error="empty_rows", hint="rows 不能为空", risk_level="medium")
    if len(rows) > _MAX_EXPORT_ROWS:
        return ToolResult(
            ok=False,
            error="too_many_rows",
            hint=f"单次导出上限 {_MAX_EXPORT_ROWS} 行",
            risk_level="medium",
        )
    out = _check_output_path(path, overwrite=overwrite, risk_level="medium")
    if isinstance(out, ToolResult):
        return out
    try:
        from openpyxl import Workbook
    except ImportError:
        return _missing_dep_result("openpyxl", "data-full", "medium")

    if all(isinstance(r, dict) for r in rows):
        header: list[str] = []
        for r in rows:
            for k in r:
                if k not in header:
                    header.append(str(k))
        data = [header, *[[r.get(k) for k in header] for r in rows]]
    elif all(isinstance(r, list) for r in rows):
        data = [[_cell_value(c) for c in r] for r in rows]
    else:
        return ToolResult(
            ok=False,
            error="mixed_row_types",
            hint="rows 必须全部为 dict 或全部为 list",
            risk_level="medium",
        )

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = (sheet_name or "Sheet1")[:31]
        for row in data:
            ws.append([str(c)[:_MAX_CELL_CHARS] if isinstance(c, str) else c for c in row])
        wb.save(str(out))
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="medium")
    return ToolResult(
        ok=True,
        content={"path": str(out), "rows_written": len(data) - 1},
        meta={"sheet": ws.title},
        risk_level="medium",
    )


def _validate_pdf_inputs(inputs: list[str], risk_level: str) -> list[Path] | ToolResult:
    if not inputs:
        return ToolResult(ok=False, error="empty_inputs", risk_level=risk_level)  # type: ignore[arg-type]
    if len(inputs) > _MAX_MERGE_INPUTS:
        return ToolResult(
            ok=False,
            error="too_many_inputs",
            hint=f"单次最多 {_MAX_MERGE_INPUTS} 个输入文件",
            risk_level=risk_level,  # type: ignore[arg-type]
        )
    paths: list[Path] = []
    for item in inputs:
        try:
            p = validate_path(item, must_exist=True)
        except Exception as exc:
            return ToolResult.from_exception(exc, risk_level=risk_level)  # type: ignore[arg-type]
        if not p.is_file() or p.suffix.lower() != ".pdf":
            return ToolResult(
                ok=False,
                error="not_a_pdf",
                hint=f"不是 PDF 文件: {p}",
                risk_level=risk_level,  # type: ignore[arg-type]
            )
        paths.append(p)
    return paths


def builtin_pdf_merge(*, inputs: list[str], output: str, overwrite: bool = False) -> ToolResult:
    """按顺序合并多个 PDF 为一个（medium 风险，走 HITL）。"""
    paths = _validate_pdf_inputs(inputs, "medium")
    if isinstance(paths, ToolResult):
        return paths
    out = _check_output_path(output, overwrite=overwrite, risk_level="medium")
    if isinstance(out, ToolResult):
        return out
    try:
        from pypdf import PdfWriter

        writer = PdfWriter()
        try:
            for p in paths:
                writer.append(str(p))
            with out.open("wb") as fh:
                writer.write(fh)
        finally:
            writer.close()
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="medium")
    return ToolResult(
        ok=True,
        content={"path": str(out), "merged_count": len(paths)},
        risk_level="medium",
    )


def _parse_page_range(spec: str, total: int) -> list[int] | None:
    """解析 "1-3,5,8-10" → 0-based 页码列表；非法返 None。"""
    pages: list[int] = []
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        m = re.fullmatch(r"(\d+)(?:-(\d+))?", part)
        if not m:
            return None
        start = int(m.group(1))
        end = int(m.group(2)) if m.group(2) else start
        if start < 1 or end > total or start > end:
            return None
        pages.extend(range(start - 1, end))
    return pages or None


def builtin_pdf_split(
    *,
    path: str,
    output_dir: str,
    pages: str | None = None,
    output_name: str | None = None,
    overwrite: bool = False,
) -> ToolResult:
    """拆分 PDF（medium 风险，走 HITL）。

    pages 为空 → 每页拆成独立文件（page_001.pdf ...）；
    pages="1-3,5" → 只抽取这些页合并写为单个文件 output_name（默认 split.pdf）。
    """
    try:
        src = validate_path(path, must_exist=True)
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="medium")
    if not src.is_file() or src.suffix.lower() != ".pdf":
        return ToolResult(ok=False, error="not_a_pdf", hint=str(src), risk_level="medium")
    try:
        out_dir = validate_path(output_dir, must_exist=False)
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="medium")
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(str(src))
        total = len(reader.pages)
        written: list[str] = []

        if pages:
            selected = _parse_page_range(pages, total)
            if selected is None:
                return ToolResult(
                    ok=False,
                    error="invalid_page_range",
                    hint=f"pages 格式应为 '1-3,5'（总页数 {total}）",
                    risk_level="medium",
                )
            name = output_name or f"{src.stem}_split.pdf"
            if not name.lower().endswith(".pdf"):
                name += ".pdf"
            target = out_dir / name
            if target.exists() and not overwrite:
                return ToolResult(
                    ok=False,
                    error="exists_no_overwrite",
                    hint=f"目标文件已存在: {target}（传 overwrite=True 覆盖）",
                    risk_level="medium",
                )
            writer = PdfWriter()
            for idx in selected:
                writer.add_page(reader.pages[idx])
            with target.open("wb") as fh:
                writer.write(fh)
            written.append(str(target))
        else:
            for idx in range(total):
                target = out_dir / f"{src.stem}_page_{idx + 1:03d}.pdf"
                if target.exists() and not overwrite:
                    return ToolResult(
                        ok=False,
                        error="exists_no_overwrite",
                        hint=f"目标文件已存在: {target}（传 overwrite=True 覆盖）",
                        risk_level="medium",
                    )
                writer = PdfWriter()
                writer.add_page(reader.pages[idx])
                with target.open("wb") as fh:
                    writer.write(fh)
                written.append(str(target))
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="medium")
    return ToolResult(
        ok=True,
        content={"files": written, "total_pages": total},
        meta={"file_count": len(written)},
        risk_level="medium",
    )


_MD_TABLE_SEP = re.compile(r"^\s*\|?\s*:?-{3,}")


def _parse_markdown_table(lines: list[str]) -> list[list[str]] | None:
    """识别 | a | b | 表格（含分隔行），返回纯文本网格；不是表格返 None。"""
    if len(lines) < 2 or "|" not in lines[0]:
        return None
    if not _MD_TABLE_SEP.match(lines[1]):
        return None
    grid: list[list[str]] = []
    for line in lines:
        if "|" not in line:
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        grid.append(cells)
    return grid if len(grid) >= 2 else None


def builtin_word_generate(
    *,
    path: str,
    markdown: str,
    title: str | None = None,
    overwrite: bool = False,
) -> ToolResult:
    """把 Markdown 子集文本生成为 Word 文档（medium 风险，走 HITL）。

    支持语法：# / ## / ### / #### 标题、- 无序列表、1. 有序列表、
    | a | b | 表格、``` 代码块（等宽字体）、普通段落。
    """
    if not markdown or not markdown.strip():
        return ToolResult(ok=False, error="empty_markdown", risk_level="medium")
    out = _check_output_path(path, overwrite=overwrite, risk_level="medium")
    if isinstance(out, ToolResult):
        return out
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        if title:
            doc.add_heading(title.strip(), level=0)

        lines = markdown.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if not stripped:
                i += 1
                continue
            # 代码块
            if stripped.startswith("```"):
                block: list[str] = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    block.append(lines[i])
                    i += 1
                para = doc.add_paragraph()
                run = para.add_run("\n".join(block))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                i += 1
                continue
            # 表格（需向后看分隔行）
            table = _parse_markdown_table(lines[i : i + 40])
            if table:
                width = max(len(r) for r in table)
                tbl = doc.add_table(rows=len(table), cols=width)
                tbl.style = "Table Grid"
                for r_idx, row in enumerate(table):
                    for c_idx in range(width):
                        tbl.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""
                # 跳过已消费的表格行（表头 + 分隔行 + 数据行）
                i += len(table)
                continue
            # 标题
            m = re.match(r"(#{1,4})\s+(.*)", stripped)
            if m:
                doc.add_heading(m.group(2), level=len(m.group(1)))
                i += 1
                continue
            # 无序列表
            if re.match(r"[-*]\s+", stripped):
                doc.add_paragraph(re.sub(r"^[-*]\s+", "", stripped), style="List Bullet")
                i += 1
                continue
            # 有序列表
            if re.match(r"\d+[.)]\s+", stripped):
                doc.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
                i += 1
                continue
            # 普通段落（合并连续非空行）
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if (
                    not nxt
                    or nxt.startswith(("#", "-", "*", "|", "```"))
                    or re.match(r"\d+[.)]\s+", nxt)
                ):
                    break
                para_lines.append(nxt)
                i += 1
            doc.add_paragraph(" ".join(para_lines))

        doc.save(str(out))
    except Exception as exc:
        return ToolResult.from_exception(exc, risk_level="medium")
    return ToolResult(
        ok=True,
        content={"path": str(out), "chars": len(markdown)},
        risk_level="medium",
    )


__all__: list[str] = [
    "builtin_excel_export",
    "builtin_excel_query",
    "builtin_pdf_merge",
    "builtin_pdf_split",
    "builtin_word_generate",
]
