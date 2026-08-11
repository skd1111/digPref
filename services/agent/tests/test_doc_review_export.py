# services/agent/tests/test_doc_review_export.py
"""doc_review · exporter（审核结果导出 Word）单元测试。"""

from __future__ import annotations

import io
import zipfile

import pytest
from agent.doc_review.exporter import build_export_docx
from docx import Document


def _parsed() -> dict:
    # 坐标语义与 parser/matcher 一致：block.start/end 与 positions 均为全文偏移
    # full_text = "甲方享有最终解释权。\n违约金为合同总额百分之百。"
    return {
        "doc_id": "d1",
        "file_name": "测试合同.docx",
        "file_path": "C:/测试合同.docx",
        "format": "docx",
        "page_count": 1,
        "full_text": "甲方享有最终解释权。\n违约金为合同总额百分之百。",
        "pages": [
            {
                "page_no": 1,
                "blocks": [
                    {"block_id": "p1b1", "text": "甲方享有最终解释权。", "start": 0, "end": 10},
                    {
                        "block_id": "p1b2",
                        "text": "违约金为合同总额百分之百。",
                        "start": 11,
                        "end": 24,
                    },
                ],
            }
        ],
    }


def _findings() -> list[dict]:
    return [
        {
            "finding_id": "f1",
            "risk_type": "compliance",
            "risk_level": "high",
            "title": "单方最终解释权",
            "description": "约定甲方享有最终解释权，属霸王条款。",
            "suggestion": "删除该条款。",
            "rule_ref": "R-001",
            "evidence_text": "甲方享有最终解释权。",
            "positions_json": '[{"page_no": 1, "block_id": "p1b1", "start": 0, "end": 10}]',
            "kb_refs": [
                {
                    "source": "《合同行政监督管理办法》",
                    "heading": "1.2 条款红线清单",
                    "excerpt": "…",
                    "matched_terms": ["解释权"],
                }
            ],
        },
        {
            "finding_id": "f2",
            "risk_type": "legal",
            "risk_level": "critical",
            "title": "违约金过高",
            "description": "违约金为合同总额百分之百，明显过高。",
            "suggestion": "",
            "rule_ref": None,
            "evidence_text": "违约金为合同总额百分之百。",
            "positions_json": '[{"page_no": 1, "block_id": "p1b2", "start": 11, "end": 24}]',
            "kb_refs": [],
        },
    ]


def test_risks_only_structured_report():
    data = build_export_docx(
        mode="risks_only",
        parsed=_parsed(),
        findings=_findings(),
        run={"doc_category": "contract", "overall_risk_level": "critical", "summary": "s"},
    )
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    # 结构要素：标题、两个风险点、原文摘录、知识库依据
    assert "风险审核报告" in text
    assert "单方最终解释权" in text
    assert "违约金过高" in text
    assert "甲方享有最终解释权。" in text
    assert "《合同行政监督管理办法》" in text
    # critical 应排在 high 之前（按等级降序）
    idx_c = text.index("违约金过高")
    idx_h = text.index("单方最终解释权")
    assert idx_c < idx_h


def test_full_mode_with_comments():
    data = build_export_docx(mode="full", parsed=_parsed(), findings=_findings())
    zf = zipfile.ZipFile(io.BytesIO(data))
    assert "word/comments.xml" in zf.namelist()
    comments_xml = zf.read("word/comments.xml").decode("utf-8")
    # 两条批注，内容含风险标题
    assert comments_xml.count("<w:comment ") == 2
    assert "单方最终解释权" in comments_xml
    assert "违约金过高" in comments_xml
    # 正文含批注锚点
    doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert doc_xml.count("commentRangeStart") == 2
    assert doc_xml.count("commentRangeEnd") == 2
    assert doc_xml.count("commentReference") == 2
    # 锚点必须包住对应原文（全文偏移→块内坐标换算正确性）：
    # 每段被批注文本的起始片段紧跟在 commentRangeStart 之后
    assert "甲方享有最终解释权" in doc_xml
    assert "违约金为合同总额百分之百" in doc_xml
    idx_anchor2 = doc_xml.index("commentRangeStart")
    idx_text1 = doc_xml.index("甲方享有最终解释权")
    assert idx_anchor2 < idx_text1
    # 关系表注册了 comments 部件
    rels_xml = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "comments.xml" in rels_xml
    # 原文仍完整
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "甲方享有最终解释权。" in text
    assert "违约金为合同总额百分之百。" in text


def test_full_mode_skips_findings_without_positions():
    findings = _findings()
    findings[0]["positions_json"] = "[]"
    data = build_export_docx(mode="full", parsed=_parsed(), findings=findings)
    zf = zipfile.ZipFile(io.BytesIO(data))
    # 只有 f2 挂批注；f1 无定位被跳过（comments 部件仅 1 条）
    comments_xml = zf.read("word/comments.xml").decode("utf-8")
    assert comments_xml.count("<w:comment ") == 1


def test_invalid_mode_raises():
    with pytest.raises(ValueError, match="unsupported export mode"):
        build_export_docx(mode="bogus", parsed=_parsed(), findings=[])
