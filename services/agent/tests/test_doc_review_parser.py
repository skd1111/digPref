# services/agent/tests/test_doc_review_parser.py
from pathlib import Path

import pytest
from agent.doc_review.parser import DocParseError, parse_document


def _write_txt(tmp_path: Path, text: str) -> Path:
    p = tmp_path / "sample.txt"
    p.write_text(text, encoding="utf-8")
    return p


def test_parse_txt_blocks_and_offsets(tmp_path):
    p = _write_txt(tmp_path, "标题\n\n第一段内容。\n\n第二段。")
    doc = parse_document(p)
    assert doc.format.value == "txt"
    assert doc.page_count == 1
    blocks = doc.pages[0].blocks
    assert [b.text for b in blocks] == ["标题", "第一段内容。", "第二段。"]
    assert blocks[1].start == len(blocks[0].text) + 1
    assert doc.full_text == "标题\n第一段内容。\n第二段。"


def test_parse_txt_gbk_fallback(tmp_path):
    p = tmp_path / "gbk.txt"
    p.write_bytes("合同金额：100万".encode("gbk"))
    doc = parse_document(p)
    assert "合同金额" in doc.full_text


def test_unsupported_format(tmp_path):
    p = tmp_path / "a.xls"
    p.write_bytes(b"x")
    with pytest.raises(DocParseError, match="不支持"):
        parse_document(p)


def test_parse_pdf(tmp_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    p = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    c.setFont("STSong-Light", 14)
    c.drawString(72, 720, "采购合同")
    c.showPage()
    c.save()
    doc = parse_document(p)
    assert doc.format.value == "pdf"
    assert doc.page_count >= 1
    assert "采购合同" in doc.full_text


def test_parse_docx(tmp_path):
    import docx

    p = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_paragraph("保密协议")
    d.add_paragraph("双方应保守商业秘密")
    d.save(str(p))
    doc = parse_document(p)
    assert doc.format.value == "docx"
    assert "保密协议" in doc.full_text
    assert "商业秘密" in doc.full_text


def test_empty_pdf_raises(tmp_path, monkeypatch):
    from agent.config import settings
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    # 关掉扫描件 OCR 回退，隔离验证“无文本层且无 OCR → 报错”（不依赖真实 OCR）
    monkeypatch.setattr(settings, "doc_review_pdf_ocr_enabled", False)
    p = tmp_path / "empty.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    c.showPage()
    c.save()
    with pytest.raises(DocParseError, match="未提取到文本"):
        parse_document(p)


def test_scanned_pdf_ocr_fallback(tmp_path, monkeypatch):
    """无文本层 PDF → 触发 OCR 回退；用 fake OCR 验证接线（不跑真实推理）。"""
    import agent.image_processing.ocr as ocr_mod
    from agent.config import settings
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    monkeypatch.setattr(settings, "doc_review_pdf_ocr_enabled", True)
    monkeypatch.setattr(ocr_mod, "rapidocr_available", lambda: True)
    monkeypatch.setattr(
        ocr_mod,
        "ocr_pdf_to_pages",
        lambda path, *, scale, max_pages: [(1, ["扫描件识别出的合同金额壹佰万"])],
    )
    p = tmp_path / "scan.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    c.showPage()  # 空白页，无文本层
    c.save()
    doc = parse_document(p)
    assert doc.format.value == "pdf"
    assert "扫描件识别出的合同金额壹佰万" in doc.full_text


def test_scanned_pdf_ocr_unavailable_raises(tmp_path, monkeypatch):
    """开关开但 RapidOCR 不可用 → 静默退化，仍报未提取到文本（不报错崩溃）。"""
    import agent.image_processing.ocr as ocr_mod
    from agent.config import settings
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    monkeypatch.setattr(settings, "doc_review_pdf_ocr_enabled", True)
    monkeypatch.setattr(ocr_mod, "rapidocr_available", lambda: False)
    p = tmp_path / "scan2.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    c.showPage()
    c.save()
    with pytest.raises(DocParseError, match="未提取到文本"):
        parse_document(p)


def test_parse_html_strips_tags(tmp_path):
    p = tmp_path / "page.html"
    p.write_text(
        "<html><head><style>x{}</style></head><body>"
        "<h1>招标公告</h1><script>bad()</script>"
        "<p>投标截止时间：2026年9月1日</p></body></html>",
        encoding="utf-8",
    )
    doc = parse_document(p)
    assert doc.format.value == "html"
    assert "招标公告" in doc.full_text
    assert "投标截止时间" in doc.full_text
    assert "bad()" not in doc.full_text
    assert "x{}" not in doc.full_text


def test_parse_csv_as_text(tmp_path):
    p = tmp_path / "data.csv"
    p.write_text("名称,金额\n合同A,100万", encoding="utf-8")
    doc = parse_document(p)
    assert doc.format.value == "txt"
    assert "合同A" in doc.full_text


def test_doc_renamed_docx_sniffed(tmp_path):
    import docx

    src = tmp_path / "real.docx"
    d = docx.Document()
    d.add_paragraph("实质是 docx 的合同")
    d.save(str(src))
    renamed = tmp_path / "fake.doc"
    renamed.write_bytes(src.read_bytes())
    doc = parse_document(renamed)
    # PK 魔数嗅探 → 按 docx 解析，format 如实报 docx
    assert doc.format.value == "docx"
    assert "实质是 docx 的合同" in doc.full_text


def test_doc_corrupt_header_raises(tmp_path):
    p = tmp_path / "broken.doc"
    p.write_bytes(b"not-ole-not-zip-content")
    with pytest.raises(DocParseError, match="文件头无法识别"):
        parse_document(p)


def test_xls_rejected_with_guidance(tmp_path):
    p = tmp_path / "old.xls"
    p.write_bytes(b"\xd0\xcf\x11\xe0fake")
    with pytest.raises(DocParseError, match=r"另存为 \.xlsx"):
        parse_document(p)
