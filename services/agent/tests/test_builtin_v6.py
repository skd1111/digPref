"""Phase 1B V6 文档处理工具族测试（2026-08-10）。

覆盖：
    - excel_query（sheets / rows / where / columns / aggregate / csv / 错误路径）
    - excel_export（list[dict] / list[list] / 覆盖保护 / 空数据）
    - pdf_merge / pdf_split（合并 / 逐页拆分 / 页码抽取 / 非法区间 / 覆盖保护）
    - word_generate（标题 / 列表 / 表格 / 代码块 / 空文本 / 覆盖保护）
    - 登记完整性：names / schema / 描述 / 风险 / registry callable
    - dispatcher 集成：read 直通 / medium 未审批 awaiting_approval / 审批后执行
"""

from __future__ import annotations

from pathlib import Path

import pytest
from agent.builtin.documents import (
    builtin_excel_export,
    builtin_excel_query,
    builtin_pdf_merge,
    builtin_pdf_split,
    builtin_word_generate,
)
from agent.builtin.models import BUILTIN_TOOL_NAMES, ToolResult
from agent.builtin.registry import TOOL_DESCRIPTIONS, TOOL_RISK_LEVEL, get_default_registry
from agent.builtin.schemas import get_builtin_schema

openpyxl = pytest.importorskip("openpyxl")


def _make_xlsx(path: Path) -> Path:
    """构造测试表格：部门 / 姓名 / 金额。"""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "明细"
    ws.append(["部门", "姓名", "金额"])
    ws.append(["销售", "张三", 100])
    ws.append(["销售", "李四", 200])
    ws.append(["研发", "王五", 50])
    wb.save(str(path))
    return path


def _make_pdf(path: Path, pages: int) -> Path:
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    with path.open("wb") as fh:
        writer.write(fh)
    return path


# ---- excel_query ---------------------------------------------------------------


class TestExcelQuery:
    @pytest.fixture
    def xlsx(self, tmp_path):
        return _make_xlsx(tmp_path / "data.xlsx")

    def test_sheets(self, xlsx):
        result = builtin_excel_query(path=str(xlsx), action="sheets")
        assert result.ok is True
        assert result.content["sheets"] == ["明细"]

    def test_rows_all(self, xlsx):
        result = builtin_excel_query(path=str(xlsx))
        assert result.ok is True
        assert result.content["header"] == ["部门", "姓名", "金额"]
        assert result.content["total"] == 3
        assert len(result.content["rows"]) == 3

    def test_rows_where_filter(self, xlsx):
        result = builtin_excel_query(path=str(xlsx), where={"部门": "销售"})
        assert result.ok is True
        assert result.content["total"] == 2

    def test_rows_columns_projection(self, xlsx):
        result = builtin_excel_query(path=str(xlsx), columns=["姓名", "金额"])
        assert result.content["header"] == ["姓名", "金额"]
        assert result.content["rows"][0] == ["张三", 100]

    def test_rows_limit_offset(self, xlsx):
        result = builtin_excel_query(path=str(xlsx), limit=1, offset=1)
        assert result.content["rows"] == [["销售", "李四", 200]]
        assert result.meta["truncated"] is True

    def test_aggregate_sum(self, xlsx):
        result = builtin_excel_query(
            path=str(xlsx), action="aggregate", group_by="部门", agg_column="金额", agg_op="sum"
        )
        assert result.ok is True
        groups = {g["部门"]: g["value"] for g in result.content["groups"]}
        assert groups == {"销售": 300, "研发": 50}

    def test_aggregate_count_no_column(self, xlsx):
        result = builtin_excel_query(
            path=str(xlsx), action="aggregate", group_by="部门", agg_op="count"
        )
        assert result.ok is True
        groups = {g["部门"]: g["count"] for g in result.content["groups"]}
        assert groups == {"销售": 2, "研发": 1}

    def test_csv_support(self, tmp_path):
        csv_file = tmp_path / "d.csv"
        csv_file.write_text("a,b\n1,2\n3,4\n", encoding="utf-8")
        result = builtin_excel_query(path=str(csv_file))
        assert result.ok is True
        assert result.content["header"] == ["a", "b"]
        assert result.content["rows"] == [["1", "2"], ["3", "4"]]

    def test_unknown_column(self, xlsx):
        result = builtin_excel_query(path=str(xlsx), where={"不存在": 1})
        assert result.ok is False
        assert result.error == "unknown_column"

    def test_invalid_action(self, xlsx):
        result = builtin_excel_query(path=str(xlsx), action="drop_table")
        assert result.ok is False
        assert result.error == "invalid_action"

    def test_invalid_agg_op(self, xlsx):
        result = builtin_excel_query(
            path=str(xlsx), action="aggregate", group_by="部门", agg_op="median"
        )
        assert result.ok is False
        assert result.error == "invalid_agg_op"

    def test_not_found(self, tmp_path):
        result = builtin_excel_query(path=str(tmp_path / "nope.xlsx"))
        assert result.ok is False


# ---- excel_export --------------------------------------------------------------


class TestExcelExport:
    def test_export_dicts(self, tmp_path):
        out = tmp_path / "out.xlsx"
        result = builtin_excel_export(
            path=str(out),
            rows=[{"姓名": "张三", "分数": 90}, {"姓名": "李四", "分数": 85, "备注": "优"}],
        )
        assert result.ok is True
        assert result.content["rows_written"] == 2
        # 回读验证
        read = builtin_excel_query(path=str(out))
        assert read.ok is True
        assert "姓名" in read.content["header"]
        assert read.content["total"] == 2

    def test_export_lists(self, tmp_path):
        out = tmp_path / "out2.xlsx"
        result = builtin_excel_export(
            path=str(out), rows=[["a", "b"], [1, 2], [3, 4]], sheet_name="数据"
        )
        assert result.ok is True
        read = builtin_excel_query(path=str(out))
        assert read.content["header"] == ["a", "b"]

    def test_empty_rows_rejected(self, tmp_path):
        result = builtin_excel_export(path=str(tmp_path / "x.xlsx"), rows=[])
        assert result.ok is False
        assert result.error == "empty_rows"

    def test_mixed_types_rejected(self, tmp_path):
        result = builtin_excel_export(path=str(tmp_path / "x.xlsx"), rows=[{"a": 1}, [1, 2]])
        assert result.ok is False
        assert result.error == "mixed_row_types"

    def test_overwrite_protection(self, tmp_path):
        out = tmp_path / "exists.xlsx"
        builtin_excel_export(path=str(out), rows=[["a"], [1]])
        result = builtin_excel_export(path=str(out), rows=[["a"], [2]])
        assert result.ok is False
        assert result.error == "exists_no_overwrite"
        result2 = builtin_excel_export(path=str(out), rows=[["a"], [2]], overwrite=True)
        assert result2.ok is True


# ---- pdf_merge / pdf_split ------------------------------------------------------


class TestPdfTools:
    @pytest.fixture
    def two_pdfs(self, tmp_path):
        a = _make_pdf(tmp_path / "a.pdf", 1)
        b = _make_pdf(tmp_path / "b.pdf", 2)
        return a, b

    def test_merge(self, tmp_path, two_pdfs):
        a, b = two_pdfs
        out = tmp_path / "merged.pdf"
        result = builtin_pdf_merge(inputs=[str(a), str(b)], output=str(out))
        assert result.ok is True
        from pypdf import PdfReader

        assert len(PdfReader(str(out)).pages) == 3

    def test_merge_non_pdf_rejected(self, tmp_path, two_pdfs):
        a, _ = two_pdfs
        txt = tmp_path / "x.txt"
        txt.write_text("hi", encoding="utf-8")
        result = builtin_pdf_merge(inputs=[str(a), str(txt)], output=str(tmp_path / "m.pdf"))
        assert result.ok is False
        assert result.error == "not_a_pdf"

    def test_merge_overwrite_protection(self, tmp_path, two_pdfs):
        a, b = two_pdfs
        out = tmp_path / "merged.pdf"
        builtin_pdf_merge(inputs=[str(a)], output=str(out))
        result = builtin_pdf_merge(inputs=[str(b)], output=str(out))
        assert result.ok is False
        assert result.error == "exists_no_overwrite"

    def test_split_all_pages(self, tmp_path, two_pdfs):
        _, b = two_pdfs
        out_dir = tmp_path / "split"
        result = builtin_pdf_split(path=str(b), output_dir=str(out_dir))
        assert result.ok is True
        assert result.content["total_pages"] == 2
        assert len(list(out_dir.glob("*.pdf"))) == 2

    def test_split_page_range(self, tmp_path):
        src = _make_pdf(tmp_path / "big.pdf", 5)
        out_dir = tmp_path / "extract"
        result = builtin_pdf_split(
            path=str(src), output_dir=str(out_dir), pages="1-2,5", output_name="part.pdf"
        )
        assert result.ok is True
        from pypdf import PdfReader

        out_file = out_dir / "part.pdf"
        assert len(PdfReader(str(out_file)).pages) == 3

    def test_split_invalid_range(self, tmp_path):
        src = _make_pdf(tmp_path / "small.pdf", 2)
        result = builtin_pdf_split(path=str(src), output_dir=str(tmp_path), pages="1-9")
        assert result.ok is False
        assert result.error == "invalid_page_range"

    def test_merge_empty_inputs(self, tmp_path):
        result = builtin_pdf_merge(inputs=[], output=str(tmp_path / "m.pdf"))
        assert result.ok is False
        assert result.error == "empty_inputs"


# ---- word_generate ---------------------------------------------------------------


class TestWordGenerate:
    MD = """# 标题一

普通段落文本。

## 小节

- 要点一
- 要点二

1. 第一步
2. 第二步

| 列A | 列B |
| --- | --- |
| 1 | 2 |

```
code_block_line
```
"""

    def test_generate_full(self, tmp_path):
        out = tmp_path / "report.docx"
        result = builtin_word_generate(path=str(out), markdown=self.MD, title="测试报告")
        assert result.ok is True
        assert out.is_file()
        from docx import Document

        doc = Document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "标题一" in text
        assert "要点一" in text
        assert "code_block_line" in text
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "列A"
        assert doc.tables[0].cell(2, 1).text == "2"

    def test_empty_markdown_rejected(self, tmp_path):
        result = builtin_word_generate(path=str(tmp_path / "x.docx"), markdown="   ")
        assert result.ok is False
        assert result.error == "empty_markdown"

    def test_overwrite_protection(self, tmp_path):
        out = tmp_path / "r.docx"
        builtin_word_generate(path=str(out), markdown="# a")
        result = builtin_word_generate(path=str(out), markdown="# b")
        assert result.ok is False
        assert result.error == "exists_no_overwrite"


# ---- 登记完整性 ----------------------------------------------------------------


_V6_TOOLS = ("excel_query", "excel_export", "pdf_merge", "pdf_split", "word_generate")


class TestRegistrationConsistency:
    def test_names_registered(self):
        for name in _V6_TOOLS:
            assert name in BUILTIN_TOOL_NAMES, f"{name} not in BUILTIN_TOOL_NAMES"

    def test_schema_description_risk(self):
        for name in _V6_TOOLS:
            assert get_builtin_schema(name) is not None, f"{name} missing schema"
            assert TOOL_DESCRIPTIONS.get(name), f"{name} missing description"
            assert name in TOOL_RISK_LEVEL, f"{name} missing risk level"

    def test_risk_levels(self):
        assert TOOL_RISK_LEVEL["excel_query"] == "read"
        for name in ("excel_export", "pdf_merge", "pdf_split", "word_generate"):
            assert TOOL_RISK_LEVEL[name] == "medium", name

    def test_registry_callable(self):
        reg = get_default_registry()
        for name in _V6_TOOLS:
            assert reg.has(name), f"{name} not registered"
            assert callable(reg.get(name))


# ---- dispatcher 集成 -------------------------------------------------------------


class TestDispatcherV6:
    async def test_excel_query_read_direct(self, tmp_path):
        """excel_query（read）不经 HITL 直接执行。"""
        from agent.builtin.dispatcher import dispatcher
        from agent.builtin.registry import reset_default_registry

        reset_default_registry()
        xlsx = _make_xlsx(tmp_path / "d.xlsx")
        result = await dispatcher().dispatch(
            {"server": "builtin", "name": "excel_query", "args": {"path": str(xlsx)}},
            {"run_id": "test-v6"},
        )
        assert result is not None
        assert result["tool_result"]["ok"] is True
        assert result["tool_result"]["content"]["total"] == 3

    async def test_word_generate_waits_for_approval(self, tmp_path, monkeypatch):
        """word_generate（medium）未审批 → HITL 前置闸门（不执行、不落盘）。"""
        from agent.builtin.dispatcher import dispatcher
        from agent.builtin.registry import reset_default_registry
        from agent.config import settings

        monkeypatch.setattr(settings, "require_hitl_for_write", True)
        reset_default_registry()
        out = tmp_path / "pending.docx"
        result = await dispatcher().dispatch(
            {
                "server": "builtin",
                "name": "word_generate",
                "args": {"path": str(out), "markdown": "# hi"},
            },
            {"run_id": "test-v6-hitl"},
        )
        assert result["awaiting_approval"] is True
        assert result["tool_result"] is None
        assert not out.exists()

    async def test_word_generate_approved_executes(self, tmp_path, monkeypatch):
        """审批通过（approval_decision=approve）后真实执行并落盘。"""
        from agent.builtin.dispatcher import dispatcher
        from agent.builtin.registry import reset_default_registry
        from agent.config import settings

        monkeypatch.setattr(settings, "require_hitl_for_write", True)
        reset_default_registry()
        out = tmp_path / "approved.docx"
        result = await dispatcher().dispatch(
            {
                "server": "builtin",
                "name": "word_generate",
                "args": {"path": str(out), "markdown": "# 已批准"},
            },
            {"run_id": "test-v6-approved", "approval_decision": "approve"},
        )
        assert result["tool_result"]["ok"] is True
        assert out.exists()

    async def test_tool_result_type(self):
        """V6 工具全部返 ToolResult（dispatcher 类型契约）。"""
        r = builtin_excel_query(path="/nonexistent/x.xlsx")
        assert isinstance(r, ToolResult)
